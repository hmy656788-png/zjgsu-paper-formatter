#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成一个模拟的未排版学术论文 .docx 文件，用于测试排版脚本。
"""

from docx import Document


def create_test_document(output_path: str = "test_input.docx"):
    """创建一个包含各种论文元素的测试文档。"""
    doc = Document()

    # 论文标题（暂不特殊处理，作为正文段落）
    doc.add_paragraph("基于多元回归模型的中国城市化进程影响因素研究")

    # 摘要
    doc.add_paragraph(
        "摘要：本文通过构建多元回归模型，从经济发展水平、产业结构、"
        "人口流动三个维度分析了中国城市化进程的主要影响因素。"
        "研究发现GDP增长率、第三产业占比和流动人口规模是推动城市化的核心动力。"
        "本研究使用2000-2020年全国31个省份的面板数据(Panel Data)进行实证分析。"
    )

    # 关键词
    doc.add_paragraph("关键词：城市化 多元回归 面板数据 影响因素 产业结构")

    # 一级标题
    doc.add_paragraph("1 引言")

    # 正文段落
    doc.add_paragraph(
        "中国的城市化进程是21世纪全球最为重要的社会经济现象之一。"
        "根据National Bureau of Statistics的数据，2020年中国城镇化率达到63.89%，"
        "较2000年的36.22%增长了近28个百分点。理解这一进程背后的驱动因素，"
        "对于制定科学的城市发展政策具有重要的理论和实践意义。"
    )

    doc.add_paragraph(
        "已有研究表明，经济增长是城市化的根本动力(Henderson, 2003)。"
        "此外，Lewis(1954)提出的二元经济理论指出，劳动力从农业部门向"
        "工业部门的转移是城市化的重要机制。"
    )

    # 一级标题
    doc.add_paragraph("2 研究设计")

    # 二级标题
    doc.add_paragraph("2.1 模型构建")

    doc.add_paragraph(
        "本文采用面板数据固定效应模型(Fixed Effects Model)进行估计。"
        "基本模型设定如下：URBit = α + β1GDPit + β2INDit + β3POPit + εit，"
        "其中URB表示城镇化率，GDP表示人均国内生产总值的对数，"
        "IND表示第三产业增加值占GDP的比重，POP表示流动人口占总人口的比例。"
    )

    # 二级标题
    doc.add_paragraph("2.2 数据来源与描述")

    doc.add_paragraph(
        "本研究使用的数据来源于《中国统计年鉴》和各省统计年鉴。"
        "样本包括2000-2020年中国31个省、自治区和直辖市的面板数据，"
        "共计651个观测值。所有经济指标均以2000年为基期进行价格平减处理。"
    )

    # 图表标题
    doc.add_paragraph("表 1 变量定义与描述性统计")

    doc.add_paragraph("（此处为表格内容占位）")

    # 一级标题
    doc.add_paragraph("3 模型的估计与检验")

    # 二级标题
    doc.add_paragraph("3.1 基准回归结果")

    doc.add_paragraph(
        "表2报告了固定效应模型的估计结果。在控制了时间固定效应后，"
        "三个核心解释变量均在1%的显著性水平下通过了t检验。"
        "具体而言，人均GDP每增长1%，城镇化率平均上升0.15个百分点。"
    )

    # 图表标题
    doc.add_paragraph("表 2 基准回归结果")

    doc.add_paragraph("（此处为表格内容占位）")

    # 二级标题
    doc.add_paragraph("3.2 稳健性检验")

    doc.add_paragraph(
        "为验证基准回归结果的稳健性，本文采用了以下策略：替换核心解释变量的度量方式、"
        "缩小样本范围（剔除直辖市）、使用工具变量法(IV)解决可能的内生性问题。"
        "结果表明核心结论保持不变。"
    )

    # 图表标题
    doc.add_paragraph("图 1 城镇化率与人均GDP散点图")

    doc.add_paragraph("（此处为图片占位）")

    # 一级标题
    doc.add_paragraph("4 结论与政策建议")

    doc.add_paragraph(
        "综合以上分析，本文得出以下主要结论：第一，经济发展水平是推动城市化的根本动力；"
        "第二，产业结构升级对城市化具有显著的正向促进作用；"
        "第三，人口流动是城市化的重要微观机制。基于上述发现，"
        "本文提出如下政策建议：持续推动经济高质量发展、"
        "加快产业结构优化升级、完善流动人口公共服务体系。"
    )

    doc.save(output_path)
    print(f"测试文档已生成：{output_path}")


if __name__ == "__main__":
    create_test_document()
