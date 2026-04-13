#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学术论文自动化排版工具 - 核心处理脚本
======================================

功能：读取未经排版的 .docx 文档，按照学术论文排版规范对其进行格式重构。

排版规则：
  1. 正文：宋体 + Times New Roman（英文/数字），小四号字，首行缩进2字符，1.5倍行距
  2. 摘要/关键词：识别并加粗标签
  3. 一级标题（如 "1 引言"）：黑体，三号，加粗，居中，段前段后1行
  4. 二级标题（如 "1.1 研究背景"）：黑体，四号，加粗，左对齐，段前段后0.5行
  5. 图表标题（如 "表 1 变量定义"）：黑体，五号，居中，无缩进

依赖安装：
  pip install python-docx

使用方法：
  python format_paper.py input.docx output.docx
"""

import re
import sys
import logging
import tempfile
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from lxml import etree

# ============================================================
# 日志配置
# ============================================================
logging.basicConfig(
    level=logging.INFO,
    format="[%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

FOOTNOTE_FONT_SIZE_PT = 10
FOOTNOTE_XML_PATH = "word/footnotes.xml"
FOOTNOTE_SKIP_TYPES = {"separator", "continuationSeparator", "continuationNotice"}
DRAWING_XML_TAGS = frozenset(
    {
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing",
        "{urn:schemas-microsoft-com:vml}shape",
        "{urn:schemas-microsoft-com:office:office}OLEObject",
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict",
    }
)
EQUATION_XML_TAGS = frozenset(
    {
        "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath",
        "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMathPara",
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object",
        "{urn:schemas-microsoft-com:office:office}OLEObject",
        "{urn:schemas-microsoft-com:vml}shape",
    }
)


def emit_progress(progress_callback, step: int, message: str, detail: str | None = None):
    """向上层报告排版进度，供 Web SSE 等场景复用。"""
    if progress_callback is None:
        return

    payload = {
        "step": step,
        "message": message,
    }
    if detail:
        payload["detail"] = detail

    try:
        progress_callback(payload)
    except Exception as exc:  # pragma: no cover - 回调失败不应影响主流程
        logger.warning(f"进度回调发送失败：{exc}")


def resolve_format_options(format_options=None) -> dict:
    """整理排版选项，兼容 bool/字符串传参并补齐默认值。"""
    resolved = DEFAULT_FORMAT_OPTIONS.copy()
    if not isinstance(format_options, dict):
        return resolved

    for key in DEFAULT_FORMAT_OPTIONS:
        value = format_options.get(key)
        if value is None:
            continue
        if isinstance(value, bool):
            resolved[key] = value
            continue
        if isinstance(value, str):
            normalized = value.strip().lower()
            if normalized in {"1", "true", "yes", "on"}:
                resolved[key] = True
                continue
            if normalized in {"0", "false", "no", "off"}:
                resolved[key] = False
                continue
        resolved[key] = bool(value)

    return resolved

# ============================================================
# 正则表达式定义（核心匹配逻辑）
# ============================================================

# --- 一级标题匹配 ---
# 匹配规则：数字或数字. + 一个或多个空格 + 至少一个中英文字符 + 标题允许字符
# 示例匹配："1 引言"、"1. 引言"、"2 研究设计"、"3 模型的估计与检验"
# 要求该段落仅包含这一行内容（独占一行），因此使用 ^ 和 $ 锚定
# 标题允许的字符包括：中英文、数字、空白、常见中英文标点符号
# 有意排除句末标点（。！？）避免把正文误判为标题
# 标题首字符：中英文字母或常见开头标点（引号、书名号等）
_HEADING_FIRST = r'[A-Za-z\u4e00-\u9fff\u201c\u2018《〈（(「『]'
_HEADING_TAIL = (
    r'[\u4e00-\u9fffA-Za-z0-9\s'
    r'\-—–~～·、，,\.:：;；/&()（）\[\]【】'
    r'%+_=*#@'
    r'\u00b7\u2013\u2014\u2026'  # · – — …
    r'\u2018\u2019\u201c\u201d'  # ' ' " "
    r'《》〈〉「」『』'           # 书名号、引号
    r']*'
)
# 负向先行断言：拒绝内部含"空格+数字"的文本（如"1 月 1 日"、"1 第 2 章"），
# 这类通常是正文里的日期/编号表达而非标题
_HEADING_NOT_INTERNAL_DIGIT = r"(?!.*\s\d)"
RE_HEADING_L1 = re.compile(
    rf"^\d+(?:\.\s*|\s+){_HEADING_NOT_INTERNAL_DIGIT}{_HEADING_FIRST}{_HEADING_TAIL}$"
)

# --- 二级标题匹配 ---
# 匹配规则：数字.数字 + 可选空格 + 至少一个中文字符
# 示例匹配："1.1研究背景"、"2.1 模型构建"、"3.2 数据来源与描述"
RE_HEADING_L2 = re.compile(
    rf"^\d+\.\d+\s*{_HEADING_NOT_INTERNAL_DIGIT}{_HEADING_FIRST}{_HEADING_TAIL}$"
)

# --- 三级标题匹配 ---
# 示例匹配："1.1.1 研究假设"、"2.3.4 稳健性检验"
RE_HEADING_L3 = re.compile(
    rf"^\d+\.\d+\.\d+\s*{_HEADING_NOT_INTERNAL_DIGIT}{_HEADING_FIRST}{_HEADING_TAIL}$"
)

# --- 图表标题匹配 ---
# 支持 "图 1 xxx"、"【图9】xxx"、"表8 xxx"、"图n"、"图1-1 xxx"、"图2.3 xxx" 等草稿写法
# index 支持章节前缀（如 1-1、2.3、1.1.1），统一归并为单独编号后重编
_CAPTION_INDEX = r"\d+(?:[\-\u2013\u2014.]\d+)*|[A-Za-z]+|[一二三四五六七八九十百千万]+"
RE_FIGURE_CAPTION = re.compile(
    rf"^(?:【\s*)?(?P<label>图)\s*(?P<index>{_CAPTION_INDEX})(?:\s*[】\]\)])?\s*[:：.\-\u2013\u2014—、]?\s*(?P<caption>.*)$"
)
RE_TABLE_CAPTION = re.compile(
    rf"^(?:【\s*)?(?P<label>表)\s*(?P<index>{_CAPTION_INDEX})(?:\s*[】\]\)])?\s*[:：.\-\u2013\u2014—、]?\s*(?P<caption>.*)$"
)
RE_EN_FIGURE_CAPTION = re.compile(
    rf"^(?P<label>figure)\s+(?P<index>{_CAPTION_INDEX})\s*[:：.\-\u2013\u2014—]\s*(?P<caption>.+)$",
    re.IGNORECASE,
)
RE_EN_TABLE_CAPTION = re.compile(
    rf"^(?P<label>table)\s+(?P<index>{_CAPTION_INDEX})\s*[:：.\-\u2013\u2014—]\s*(?P<caption>.+)$",
    re.IGNORECASE,
)
RE_EN_TABLE_LABEL_ONLY = re.compile(r"^(?P<label>table)\s+(?P<index>\d+(?:\.\d+)*)\s*$", re.IGNORECASE)
RE_EN_FIGURE_LABEL_ONLY = re.compile(r"^(?P<label>figure)\s+(?P<index>\d+(?:\.\d+)*)\s*$", re.IGNORECASE)

# --- 摘要标识匹配 ---
# 匹配规则：段落起始处包含 "摘要" + 可选的标点符号（如 ":"、"："）
RE_ABSTRACT = re.compile(r"^摘\s*要\s*[:：]?\s*")

# --- 关键词标识匹配 ---
# 匹配规则：段落起始处包含 "关键词" + 可选的标点符号
RE_KEYWORDS = re.compile(r"^关\s*键\s*词\s*[:：]?\s*")
RE_ENGLISH_ABSTRACT_HEADING = re.compile(r"^(?:英文摘要|abstract)\s*$", re.IGNORECASE)
RE_ENGLISH_ABSTRACT = re.compile(r"^abstract\s*[:：]\s*", re.IGNORECASE)
RE_ENGLISH_KEYWORDS = re.compile(r"^(?:keywords?|key\s*words?)\s*[:：]?\s*", re.IGNORECASE)
RE_REFERENCES_HEADING = re.compile(r"^(?:参\s*考\s*文\s*献|references?)\s*[:：]?\s*$", re.IGNORECASE)
RE_CAPTION_NOTE = re.compile(r"^(?:注|说明|资料来源|数据来源|来源|source|note)\s*[:：.]?\s*", re.IGNORECASE)
RE_SECTION_HEADING = re.compile(
    r"^(?:"
    r"致\s*谢"
    r"|附\s*录(?:\s*[A-Za-z一二三四五六七八九十\d]{1,3})?"
    r"|作\s*者\s*简\s*介"
    r"|基\s*金\s*项\s*目"
    r"|英\s*文\s*摘\s*要"
    r"|abstract"
    r"|acknowledg(?:e)?ments?"
    r")\s*$",
    re.IGNORECASE,
)
RE_REFERENCE_ENTRY_TEXT = re.compile(
    r"^(?:\[\d+\]|\(\d+\)|（\d+）|\d+\.\s+|\d+、\s*).+"
)
RE_EN_REFERENCE_ENTRY_TEXT = re.compile(
    r"^[A-Z][A-Za-z'`-]+(?:,\s+[A-Z](?:\.[A-Z])?\.?)+(?:,\s*&\s*[A-Z][A-Za-z'`-]+(?:,\s+[A-Z](?:\.[A-Z])?\.?)+)*\s+\(\d{4}[a-z]?\)\.",
)
RE_TITLE_METADATA_PREFIX = re.compile(
    r"^(作者|姓名|学院|学校|专业|指导教师|导师|学号|班级|单位|联系方式|电话|邮箱|电子邮箱|email|e-mail)\s*[:：]",
    re.IGNORECASE,
)
RE_BASIC_NUMBERED_HEADING = re.compile(r"^\d+(?:\.\d+){0,2}(?:\.)?\s+\S.+$")
RE_CJK_CHAR = re.compile(r"[\u4e00-\u9fff]")
RE_LATIN_CHAR = re.compile(r"[A-Za-z]")

PAGE_LAYOUT = {
    "page_size": "A4",
    "page_width_cm": 21.0,
    "page_height_cm": 29.7,
    "margins_cm": {
        "top": 2.54,
        "bottom": 2.54,
        "left": 3.18,
        "right": 3.18,
    },
    "header_distance_cm": 1.5,
    "footer_distance_cm": 1.5,
}
DEFAULT_HEADER_TEXT = ""
RUNNING_HEADER_MAX_LENGTH = 28
CAPTION_MAX_LENGTH = 60
EN_CAPTION_MAX_LENGTH = 120
ENGLISH_TEMPLATE_BODY_FIRST_INDENT_PT = 36
ENGLISH_TEMPLATE_REFERENCE_HANGING_PT = 36
MIN_TOC_HEADING_COUNT = 2
STYLED_HEADING_MAX_LENGTH = 120
COVER_INFO_FIELDS = [
    ("学院", "college"),
    ("教师", "teacher"),
    ("班级", "class_name"),
    ("姓名", "student_name"),
    ("学号", "student_id"),
]
COVER_LAYOUT = {
    "logo_width_cm": 3.6,
    "school_name_width_cm": 10.6,
    "logo_space_before_pt": 30,
    "logo_space_after_pt": 10,
    "school_name_space_after_pt": 42,
    "title_size_pt": 26,
    "title_space_after_pt": 90,
    "info_spacer_after_pt": 16,
    "info_font_pt": 15,
    "label_width_cm": 2.2,
    "value_width_cm": 7.2,
}
COVER_IMAGE_CANDIDATES = {
    "logo": ("logo.png", "logo.jpg", "logo.jpeg"),
    "school_name": (
        "school_name.png",
        "school_name.jpg",
        "school_name.jpeg",
        "school_name_calligraphy.png",
        "school_name_calligraphy.jpg",
        "浙江工商大学.png",
        "浙江工商大学.jpg",
    ),
}
DEFAULT_FORMAT_OPTIONS = {
    "insert_toc": True,
    "resize_images": True,
    "format_footnotes": True,
}


# ============================================================
# 段落分类枚举
# ============================================================
class ParagraphType:
    """段落类型常量"""
    TITLE = "title"                  # 论文标题
    HEADING_L1 = "heading_l1"        # 一级标题
    HEADING_L2 = "heading_l2"        # 二级标题
    HEADING_L3 = "heading_l3"        # 三级标题
    FIGURE_CAPTION = "figure_caption"  # 图标题
    TABLE_CAPTION = "table_caption"    # 表标题
    CAPTION_NOTE = "caption_note"      # 图表附注/来源说明
    SECTION_HEADING = "section_heading"  # 非编号章节标题（如致谢/附录）
    REFERENCES_HEADING = "references_heading"  # 参考文献标题
    REFERENCE_ENTRY = "reference_entry"        # 参考文献条目
    ABSTRACT = "abstract"            # 摘要段落
    KEYWORDS = "keywords"            # 关键词段落
    ENGLISH_ABSTRACT_HEADING = "english_abstract_heading"  # 英文摘要标题
    ENGLISH_ABSTRACT = "english_abstract"                  # 英文摘要正文
    ENGLISH_KEYWORDS = "english_keywords"                  # 英文关键词
    BODY = "body"                    # 正文段落


@dataclass(slots=True)
class ParagraphAnalysis:
    """缓存单个段落的预分析结果，避免主流程重复做文本和 XML 扫描。"""

    index: int
    normalized_text: str
    classified_type: str
    caption_match: tuple[str, re.Match[str], str] | None
    inferred_heading_type: str | None
    has_drawing: bool
    has_equation: bool
    is_reference_entry_candidate: bool
    is_caption_note_candidate: bool


HEADING_LEVEL_BY_TYPE = {
    ParagraphType.HEADING_L1: 0,
    ParagraphType.HEADING_L2: 1,
    ParagraphType.HEADING_L3: 2,
}
HEADING_NUMBER_PATTERNS = {
    ParagraphType.HEADING_L1: re.compile(r"^(?P<number>\d+)(?:\.\s*|\s+)(?P<title>.+)$"),
    ParagraphType.HEADING_L2: re.compile(r"^(?P<number>\d+\.\d+)\s*(?P<title>.+)$"),
    ParagraphType.HEADING_L3: re.compile(r"^(?P<number>\d+\.\d+\.\d+)\s*(?P<title>.+)$"),
}
HEADING_NUMBERING_NSID = "5A475355"
HEADING_NUMBERING_TEMPLATE = "5A475355"
HEADING_NUMBERING_LEVEL_TEXTS = ("%1", "%1.%2", "%1.%2.%3")
UNNUMBERED_HEADING_MAX_LENGTH = 40
HEADING_STYLE_NAME_TO_LEVEL = {
    "heading1": 0,
    "heading2": 1,
    "heading3": 2,
    "标题1": 0,
    "标题2": 1,
    "标题3": 2,
}


# 手动输入的图目录/表目录条目特征：行末含"引导点 + 页码"
# 如 "图 1 xxx......3"、"表 2 xxx……5"
# 规范化文本里连续空格已被折叠，故只检测点/省略号等非空白引导符
RE_TOC_ENTRY_TAIL = re.compile(
    r"(?:\.{3,}|…+|\u3002{3,}|·{3,})\s*\d+\s*$"
)


def looks_like_toc_entry(normalized_text: str) -> bool:
    """判断段落是否像手动输入的目录条目（引导点 + 页码结尾）。"""
    if not normalized_text:
        return False
    return bool(RE_TOC_ENTRY_TAIL.search(normalized_text))


# ============================================================
# 段落分类函数
# ============================================================
def match_caption_in_normalized_text(normalized_text: str):
    """识别已规范化文本中的图/表标题。"""
    normalized = (normalized_text or "").strip()
    if not normalized:
        return None

    # 目录条目（如 "图 1 xxx......3"）尾部含引导点+页码，不视为图表标题
    if looks_like_toc_entry(normalized):
        return None

    if len(normalized) <= CAPTION_MAX_LENGTH:
        figure_match = RE_FIGURE_CAPTION.match(normalized)
        if figure_match:
            return ParagraphType.FIGURE_CAPTION, figure_match, normalized

        table_match = RE_TABLE_CAPTION.match(normalized)
        if table_match:
            return ParagraphType.TABLE_CAPTION, table_match, normalized

    if len(normalized) <= EN_CAPTION_MAX_LENGTH:
        en_figure_match = RE_EN_FIGURE_CAPTION.match(normalized)
        if en_figure_match:
            return ParagraphType.FIGURE_CAPTION, en_figure_match, normalized

        en_table_match = RE_EN_TABLE_CAPTION.match(normalized)
        if en_table_match:
            return ParagraphType.TABLE_CAPTION, en_table_match, normalized

    return None


def classify_normalized_paragraph(
    normalized_text: str,
    *,
    caption_match: tuple[str, re.Match[str], str] | None = None,
) -> str:
    """
    根据已规范化的段落文本判断其类型。

    分类优先级（从高到低）：
      1. 摘要 → 包含"摘要："开头
      2. 关键词 → 包含"关键词："开头
      3. 参考文献标题 → "参考文献"
      4. 非编号章节标题 → "致谢"、"附录" 等独占标题
      5. 一级标题 → "数字 空格 中英文" 格式
      6. 二级标题 → "数字.数字 中英文" 格式
      7. 三级标题 → "数字.数字.数字 中英文" 格式
      8. 图标题 / 表标题 → "图 1"、"【图9】"、"表8" 等短标题
      9. 正文 → 以上都不匹配时的默认类型

    Args:
        normalized_text: 已经做过空白规范化的段落文本

    Returns:
        ParagraphType 常量字符串
    """
    stripped = (normalized_text or "").strip()

    if not stripped:
        return ParagraphType.BODY  # 空段落当作正文处理

    # 优先匹配摘要和关键词
    if RE_ENGLISH_ABSTRACT_HEADING.match(stripped):
        return ParagraphType.ENGLISH_ABSTRACT_HEADING

    if RE_ABSTRACT.match(stripped):
        return ParagraphType.ABSTRACT

    if RE_KEYWORDS.match(stripped):
        return ParagraphType.KEYWORDS

    if RE_ENGLISH_ABSTRACT.match(stripped):
        return ParagraphType.ENGLISH_ABSTRACT

    if RE_ENGLISH_KEYWORDS.match(stripped):
        return ParagraphType.ENGLISH_KEYWORDS

    if RE_REFERENCES_HEADING.match(stripped):
        return ParagraphType.REFERENCES_HEADING

    if RE_SECTION_HEADING.match(stripped):
        return ParagraphType.SECTION_HEADING

    # 目录条目（如 "1 引言......3"）含引导点+页码，不是真正的标题
    if looks_like_toc_entry(stripped):
        return ParagraphType.BODY

    # 匹配一级标题（注意：先匹配一级，再匹配二级，避免误判）
    if RE_HEADING_L1.match(stripped):
        return ParagraphType.HEADING_L1

    # 匹配二级标题
    if RE_HEADING_L2.match(stripped):
        return ParagraphType.HEADING_L2

    # 匹配三级标题
    if RE_HEADING_L3.match(stripped):
        return ParagraphType.HEADING_L3

    if caption_match is None:
        caption_match = match_caption_in_normalized_text(stripped)
    if caption_match:
        return caption_match[0]

    # 默认为正文
    return ParagraphType.BODY


def classify_paragraph(text: str) -> str:
    """根据段落纯文本内容判断其类型。"""
    normalized_text = normalize_text_for_matching(text)
    return classify_normalized_paragraph(normalized_text)


def split_text_to_paragraphs(text: str) -> list[str]:
    """
    规范化纯文本输入并拆分为段落列表。

    - 统一处理 Windows/macOS/Linux 换行符
    - 保留中间空行，便于在导出的文档中保留段落间距
    - 避免把换行控制字符残留到段落正文里
    """
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n").lstrip("\ufeff")
    lines = normalized.split("\n")

    while lines and lines[-1] == "":
        lines.pop()

    return lines or [""]


def normalize_text_for_matching(text: str) -> str:
    """
    规范化段落文本，便于处理软回车、全角空格和多个连续空白。

    对标题、图表标题等“独占一行”的识别尤其重要。
    """
    normalized = (text or "").replace("\r", "\n").replace("\v", "\n").replace("\u00a0", " ").replace("\u3000", " ")
    normalized = re.sub(r"\n+", " ", normalized)
    normalized = re.sub(r"[ \t]+", " ", normalized)
    return normalized.strip()


def match_caption(text: str):
    """
    识别图/表标题，并返回类型、匹配结果和规范化后的文本。

    通过长度限制降低误把正文识别为图表标题的风险。
    """
    normalized = normalize_text_for_matching(text)
    return match_caption_in_normalized_text(normalized)


def rebuild_caption_text(kind: str, number: int, match: re.Match) -> str:
    """按照统一编号规则重建图/表标题文本。"""
    raw_label = (match.groupdict().get("label") or "").strip()
    if raw_label:
        if raw_label.lower() == "figure":
            label = "Figure"
        elif raw_label.lower() == "table":
            label = "Table"
        else:
            label = raw_label
    else:
        label = "图" if kind == ParagraphType.FIGURE_CAPTION else "表"
    caption = (match.group("caption") or "").strip()
    return f"{label} {number}" if not caption else f"{label} {number} {caption}"


def is_caption_note_candidate(text: str, normalized_text: str | None = None) -> bool:
    """判断段落是否像图表后的附注/来源说明。"""
    normalized = normalized_text if normalized_text is not None else normalize_text_for_matching(text)
    return bool(RE_CAPTION_NOTE.match(normalized))


def is_title_candidate(
    text: str,
    *,
    normalized_text: str | None = None,
    classified_type: str | None = None,
) -> bool:
    """判断一个段落是否像论文标题。"""
    stripped = normalized_text if normalized_text is not None else normalize_text_for_matching(text)

    if not stripped:
        return False

    if not 6 <= len(stripped) <= 40:
        return False

    if stripped.endswith(("。", "！", "？", "!", "?", "；", ";")):
        return False

    if "@" in stripped or RE_TITLE_METADATA_PREFIX.match(stripped):
        return False

    para_type = classified_type if classified_type is not None else classify_normalized_paragraph(stripped)
    return para_type == ParagraphType.BODY


def is_reference_entry_text(text: str, normalized_text: str | None = None) -> bool:
    """判断参考文献段落是否具备常见的编号前缀。"""
    normalized = normalized_text if normalized_text is not None else normalize_text_for_matching(text)
    return bool(RE_REFERENCE_ENTRY_TEXT.match(normalized) or RE_EN_REFERENCE_ENTRY_TEXT.match(normalized))


def find_title_paragraph_index(paragraphs, analyses: list[ParagraphAnalysis] | None = None) -> int | None:
    """
    尝试识别论文主标题。

    规则保持保守：
    - 只考虑前 3 个非空段落中的候选项
    - 段落本身必须像标题
    - 后续 3 个非空段落内需要出现“摘要”或“关键词”
    """
    if analyses is None:
        non_empty = []
        for index, paragraph in enumerate(paragraphs):
            normalized_text = normalize_text_for_matching(paragraph.text)
            if not normalized_text:
                continue
            non_empty.append(
                (
                    index,
                    normalized_text,
                    classify_normalized_paragraph(normalized_text),
                )
            )
    else:
        non_empty = [
            (analysis.index, analysis.normalized_text, analysis.classified_type)
            for analysis in analyses
            if analysis.normalized_text
        ]

    if not non_empty:
        return None

    for candidate_position, (candidate_index, candidate_text, candidate_type) in enumerate(non_empty[:3]):
        if not is_title_candidate(
            candidate_text,
            normalized_text=candidate_text,
            classified_type=candidate_type,
        ):
            continue

        for _, text, para_type in non_empty[candidate_position + 1:candidate_position + 4]:
            if para_type in {
                ParagraphType.ABSTRACT,
                ParagraphType.KEYWORDS,
                ParagraphType.ENGLISH_ABSTRACT_HEADING,
                ParagraphType.ENGLISH_ABSTRACT,
                ParagraphType.ENGLISH_KEYWORDS,
            }:
                return candidate_index

            if para_type in {
                ParagraphType.HEADING_L1,
                ParagraphType.HEADING_L2,
                ParagraphType.FIGURE_CAPTION,
                ParagraphType.TABLE_CAPTION,
            }:
                break

    return None


def extract_heading_numbering(text: str, para_type: str) -> tuple[str, tuple[int, ...]]:
    """
    从标题文本中提取编号前缀和纯标题文本。

    例如：
      - "1 引言"    -> ("引言", (1,))
      - "1.1 背景"  -> ("背景", (1, 1))
      - "1.1.1 假设" -> ("假设", (1, 1, 1))
    """
    normalized = normalize_text_for_matching(text)
    pattern = HEADING_NUMBER_PATTERNS.get(para_type)
    if pattern is None:
        return normalized, ()

    match = pattern.match(normalized)
    if match is None:
        return normalized, ()

    number_text = (match.group("number") or "").rstrip(".")
    title_text = normalize_text_for_matching(match.group("title"))
    if not number_text or not title_text:
        return normalized, ()

    return title_text, tuple(int(part) for part in number_text.split("."))


def _read_outline_level(p_pr) -> int | None:
    """从段落或样式的 pPr 中读取 outline level。"""
    if p_pr is None:
        return None

    outline_level = p_pr.find(qn("w:outlineLvl"))
    if outline_level is None:
        return None

    raw_value = outline_level.get(qn("w:val"))
    if raw_value is None:
        return None

    try:
        return int(raw_value)
    except (TypeError, ValueError):
        return None


def _get_paragraph_outline_level_hint(paragraph) -> int | None:
    """
    获取段落自带的标题层级提示。

    优先读取段落直接设置的 outline level；如果没有，再回退到样式中配置的
    outline level 或常见的 Heading 样式名称。
    """
    direct_level = _read_outline_level(paragraph._element.pPr)
    if direct_level in {0, 1, 2}:
        return direct_level

    style = paragraph.style
    if style is None:
        return None

    style_level = _read_outline_level(style.element.find(qn("w:pPr")))
    if style_level in {0, 1, 2}:
        return style_level

    style_name = re.sub(r"\s+", "", str(getattr(style, "name", "") or ""))
    return HEADING_STYLE_NAME_TO_LEVEL.get(style_name.lower(), HEADING_STYLE_NAME_TO_LEVEL.get(style_name))


def looks_like_unnumbered_heading(text: str, normalized_text: str | None = None) -> bool:
    """保守判断一段文字是否像“缺失编号的标题”。"""
    normalized = normalized_text if normalized_text is not None else normalize_text_for_matching(text)
    if not normalized or len(normalized) > UNNUMBERED_HEADING_MAX_LENGTH:
        return False

    if normalized.endswith(("。", "！", "？", "!", "?", "；", ";", "，", ",", "：", ":")):
        return False

    # 目录条目不应被推断为无编号标题
    if looks_like_toc_entry(normalized):
        return False

    if is_reference_entry_text(normalized, normalized_text=normalized):
        return False

    if match_caption_in_normalized_text(normalized):
        return False

    return True


def looks_like_styled_heading(text: str, normalized_text: str | None = None) -> bool:
    """判断带样式提示的段落是否仍可视为标题。"""
    normalized = normalized_text if normalized_text is not None else normalize_text_for_matching(text)
    if not normalized or len(normalized) > STYLED_HEADING_MAX_LENGTH:
        return False

    if looks_like_toc_entry(normalized):
        return False

    if is_reference_entry_text(normalized, normalized_text=normalized):
        return False

    if match_caption_in_normalized_text(normalized):
        return False

    if normalized.endswith(("。", "！", "？", "!", "?", "；", ";", "，", ",")):
        return False

    if RE_BASIC_NUMBERED_HEADING.match(normalized):
        return True

    return looks_like_unnumbered_heading(text, normalized_text=normalized)


def infer_heading_type_from_paragraph(
    paragraph,
    text: str,
    normalized_text: str | None = None,
) -> str | None:
    """根据段落原始样式/outline level，推断标题层级。"""
    normalized = normalized_text if normalized_text is not None else normalize_text_for_matching(text)
    level = _get_paragraph_outline_level_hint(paragraph)
    if level not in {0, 1, 2}:
        return None

    if not looks_like_styled_heading(text, normalized_text=normalized):
        return None

    return {
        0: ParagraphType.HEADING_L1,
        1: ParagraphType.HEADING_L2,
        2: ParagraphType.HEADING_L3,
    }.get(level)


def is_english_dominant_text(text: str) -> bool:
    """粗略判断一段文字是否以英文为主，便于调整英文文档的版式。"""
    normalized = normalize_text_for_matching(text)
    if not normalized or RE_CJK_CHAR.search(normalized):
        return False

    latin_count = len(RE_LATIN_CHAR.findall(normalized))
    if latin_count < 6:
        return False

    non_space = re.sub(r"\s+", "", normalized)
    if not non_space:
        return False

    return latin_count / len(non_space) >= 0.45


def detect_english_template_mode(paragraphs, analyses: list[ParagraphAnalysis]) -> bool:
    """判断当前文档是否更接近用户提供的英文论文模板。"""
    abstract_index = next(
        (analysis.index for analysis in analyses if analysis.classified_type == ParagraphType.ENGLISH_ABSTRACT_HEADING),
        None,
    )
    if abstract_index is None or abstract_index < 4:
        return False

    front_matter = [
        (paragraph, analysis)
        for paragraph, analysis in zip(paragraphs[:abstract_index], analyses[:abstract_index])
        if analysis.normalized_text
    ]
    if len(front_matter) < 4:
        return False

    centered_english_front_count = sum(
        1
        for paragraph, analysis in front_matter
        if is_english_dominant_text(analysis.normalized_text)
        and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    )
    if centered_english_front_count < 3:
        return False

    english_heading_count = sum(
        1
        for paragraph, analysis in zip(paragraphs, analyses)
        if analysis.normalized_text
        and _get_paragraph_outline_level_hint(paragraph) in {0, 1, 2}
        and is_english_dominant_text(analysis.normalized_text)
    )
    if english_heading_count < 4:
        return False

    return any(
        analysis.classified_type == ParagraphType.REFERENCES_HEADING
        and is_english_dominant_text(analysis.normalized_text)
        for analysis in analyses
    )


def find_english_front_matter_roles(analyses: list[ParagraphAnalysis]) -> dict[int, str]:
    """识别英文模板中摘要前的居中封面/署名区。"""
    first_structural_index = next(
        (
            analysis.index
            for analysis in analyses
            if analysis.classified_type in {
                ParagraphType.ENGLISH_ABSTRACT_HEADING,
                ParagraphType.ABSTRACT,
                ParagraphType.HEADING_L1,
                ParagraphType.HEADING_L2,
                ParagraphType.HEADING_L3,
                ParagraphType.REFERENCES_HEADING,
                ParagraphType.SECTION_HEADING,
            }
        ),
        None,
    )
    if first_structural_index in {None, 0}:
        return {}

    roles: dict[int, str] = {}
    groups: list[list[int]] = []
    current_group: list[int] = []

    for index in range(first_structural_index):
        normalized = analyses[index].normalized_text
        if normalized:
            current_group.append(index)
            continue
        roles[index] = "spacer"
        if current_group:
            groups.append(current_group)
            current_group = []

    if current_group:
        groups.append(current_group)

    if not groups:
        return roles

    for index in groups[0]:
        roles[index] = "title"
    for group in groups[1:]:
        for index in group:
            roles[index] = "meta"

    return roles


def looks_like_english_caption_title(text: str) -> bool:
    """判断是否像英文模板里表题编号后的说明行。"""
    normalized = normalize_text_for_matching(text)
    if not normalized or len(normalized) > EN_CAPTION_MAX_LENGTH:
        return False

    if not is_english_dominant_text(normalized):
        return False

    if RE_CAPTION_NOTE.match(normalized):
        return False

    if classify_normalized_paragraph(normalized) != ParagraphType.BODY:
        return False

    return not normalized.endswith((".", "!", "?", ":", ";"))


def find_english_split_caption_roles(analyses: list[ParagraphAnalysis]) -> dict[int, str]:
    """识别英文模板中拆成两行的 Table/Figure 标题。"""
    roles: dict[int, str] = {}

    for analysis in analyses:
        normalized = analysis.normalized_text
        if not normalized:
            continue

        if RE_EN_TABLE_LABEL_ONLY.match(normalized):
            next_index = analysis.index + 1
            if next_index < len(analyses) and looks_like_english_caption_title(analyses[next_index].normalized_text):
                roles[analysis.index] = "table_label"
                roles[next_index] = "table_title"

        if RE_EN_FIGURE_LABEL_ONLY.match(normalized):
            next_index = analysis.index + 1
            if next_index < len(analyses) and looks_like_english_caption_title(analyses[next_index].normalized_text):
                roles[analysis.index] = "figure_label"
                roles[next_index] = "figure_title"

    return roles


def should_preserve_explicit_heading_text(text: str, para_type: str, english_template_mode: bool) -> bool:
    """保留显式编号文本，避免 WPS 对多级编号的渲染出现补点或层级错位。"""
    pattern = HEADING_NUMBER_PATTERNS.get(para_type)
    if pattern is None:
        return False

    normalized = normalize_text_for_matching(text)
    if pattern.match(normalized) is None:
        return False

    return english_template_mode or not is_english_dominant_text(normalized)


def rebuild_explicit_heading_text(heading_text: str, numbering_parts: tuple[int, ...]) -> str:
    """把显式编号标题规范化为纯文本，统一一级标题后的句点与空格。"""
    if not numbering_parts:
        return heading_text

    number_text = ".".join(str(part) for part in numbering_parts)
    return f"{number_text} {heading_text}".strip()


def _build_paragraph_analyses(paragraphs) -> list[ParagraphAnalysis]:
    """预分析所有正文段落，复用规范化文本、分类结果与公式探测。"""
    analyses = []
    for index, paragraph in enumerate(paragraphs):
        raw_text = paragraph.text
        normalized_text = normalize_text_for_matching(raw_text)
        caption_match = match_caption_in_normalized_text(normalized_text)
        classified_type = classify_normalized_paragraph(
            normalized_text,
            caption_match=caption_match,
        )
        has_drawing, has_equation = _scan_paragraph_content(paragraph)
        inferred_heading_type = None
        if classified_type == ParagraphType.BODY:
            inferred_heading_type = infer_heading_type_from_paragraph(
                paragraph,
                raw_text,
                normalized_text=normalized_text,
            )

        analyses.append(
            ParagraphAnalysis(
                index=index,
                normalized_text=normalized_text,
                classified_type=classified_type,
                caption_match=caption_match,
                inferred_heading_type=inferred_heading_type,
                has_drawing=has_drawing,
                has_equation=has_equation,
                is_reference_entry_candidate=is_reference_entry_text(
                    normalized_text,
                    normalized_text=normalized_text,
                ),
                is_caption_note_candidate=is_caption_note_candidate(
                    normalized_text,
                    normalized_text=normalized_text,
                ),
            )
        )

    return analyses


# TOC 相关的样式名称（小写、去空格后匹配）
_TOC_STYLE_NAMES = {
    "toc1", "toc2", "toc3", "toc4", "toc5", "toc6", "toc7", "toc8", "toc9",
    "tableoffigures", "listoffigures", "listoftables",
    "目录1", "目录2", "目录3", "目录4", "目录5",
}


def _find_field_paragraph_indices(paragraphs) -> set[int]:
    """
    检测处于 Word 域（field）块内部的段落索引。

    Word 中的目录（TOC）、图目录、表目录等使用域代码实现。域代码由
    fldChar begin … fldChar separate … fldChar end 三段组成，
    其中 separate 和 end 之间的段落是域的渲染内容（即目录条目）。
    这些段落的文本与图表标题相似，但不应被重新格式化。

    同时检测段落样式名包含 "TOC"、"目录"、"Table of Figures" 等
    的情况作为补充判断。
    """
    field_para_indices: set[int] = set()
    field_depth = 0

    for idx, paragraph in enumerate(paragraphs):
        # 检查样式名
        style = paragraph.style
        if style is not None:
            style_name = re.sub(r"\s+", "", str(getattr(style, "name", "") or "")).lower()
            if style_name in _TOC_STYLE_NAMES:
                field_para_indices.add(idx)

        # 检查是否含有指向 TOC 书签的超链接（Word 自动生成的目录条目特征）
        # 只检测 _Toc 前缀；_Ref 是正文交叉引用（如"见图1"），不应跳过
        for hyperlink in paragraph._element.iter(qn("w:hyperlink")):
            anchor = hyperlink.get(qn("w:anchor"), "") or ""
            if anchor.startswith("_Toc"):
                field_para_indices.add(idx)
                break

        # 记录进入段落时的域深度；如果已经在域内，本段落必然是域内容
        in_field = field_depth > 0

        # 递归查找段落内所有 fldChar，按文档顺序追踪域深度
        # 用 iter() 覆盖嵌套在 w:hyperlink、w:sdt、w:smartTag 等容器内的情况
        for fld_char in paragraph._element.iter(qn("w:fldChar")):
            fld_type = fld_char.get(qn("w:fldCharType"))
            if fld_type == "begin":
                field_depth += 1
                in_field = True
            elif fld_type == "end":
                field_depth = max(0, field_depth - 1)
                # 结束符所在段落本身也属于域的一部分
                in_field = True
            elif fld_type == "separate":
                in_field = True

        # 处理完段落后，若当前仍在域内，则仍属于域内容
        if field_depth > 0:
            in_field = True

        if in_field:
            field_para_indices.add(idx)

    return field_para_indices


def _log_detected_paragraph(label: str, index: int, text: str) -> None:
    """逐段调试日志仅在 DEBUG 下输出，避免大文档时产生过多日志 I/O。"""
    if not logger.isEnabledFor(logging.DEBUG):
        return

    preview = text if len(text) <= 30 else f"{text[:30]}..."
    logger.debug(f'  [{label}] 第{index + 1}段: "{preview}"')


def resolve_heading_numbering_parts(
    para_type: str,
    explicit_parts: tuple[int, ...],
    numbering_state: list[int],
    *,
    allow_auto_numbering: bool = False,
) -> tuple[int, ...]:
    """综合显式编号和推断层级，返回当前标题应使用的原生编号。"""
    level = HEADING_LEVEL_BY_TYPE.get(para_type)
    if level is None:
        return ()

    if explicit_parts:
        for index in range(len(numbering_state)):
            numbering_state[index] = explicit_parts[index] if index < len(explicit_parts) else 0
        return explicit_parts

    if not allow_auto_numbering:
        return ()

    if level > 0 and any(numbering_state[index] <= 0 for index in range(level)):
        return ()

    numbering_state[level] = numbering_state[level] + 1 if numbering_state[level] > 0 else 1
    for index in range(level + 1, len(numbering_state)):
        numbering_state[index] = 0

    return tuple(numbering_state[: level + 1])


# ============================================================
# 底层格式设置工具函数
# ============================================================
def _set_run_font(run, cn_font: str, en_font: str, size_pt: float, bold: bool | None = None, color: RGBColor = None):
    """
    设置 run 级别的字体属性。

    通过直接操作底层 XML 确保中文字体（eastAsia）和西文字体分别正确设置。
    python-docx 的高级 API 无法单独设置 eastAsia 字体，因此需要手动操作 XML。

    Args:
        run:      docx Run 对象
        cn_font:  中文字体名称（如 "宋体"、"黑体"）
        en_font:  西文字体名称（如 "Times New Roman"）
        size_pt:  字号磅值
        bold:     是否加粗；传 None 时保留原有加粗状态
        color:    字体颜色（可选）
    """
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    run.font.name = en_font  # 设置西文（Latin）字体

    if color:
        run.font.color.rgb = color

    # 通过底层 XML 设置中文字体（eastAsia）
    # python-docx 不直接支持 eastAsia 字体设置，需要操作 XML
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        r_pr.insert(0, r_fonts)

    r_fonts.set(qn("w:eastAsia"), cn_font)
    r_fonts.set(qn("w:ascii"), en_font)
    r_fonts.set(qn("w:hAnsi"), en_font)
    r_fonts.set(qn("w:cs"), en_font)  # 复杂脚本字体也设为西文字体


def _remove_all_runs(paragraph):
    """删除段落中已有的 run，便于重建页眉页脚等结构。"""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _replace_paragraph_text(paragraph, text: str):
    """安全重建纯文本段落内容。仅用于标题、图表标题等纯文本段落。"""
    _remove_all_runs(paragraph)
    if text:
        paragraph.add_run(text)


def _apply_run_fonts(paragraph, cn_font: str, en_font: str, size_pt: float, bold: bool | None = None):
    """遍历段落中所有 run，统一设置中西文字体。"""
    for run in paragraph.runs:
        current_bold = run.font.bold if bold is None else bold
        _set_run_font(run, cn_font=cn_font, en_font=en_font, size_pt=size_pt, bold=current_bold)


def _is_list_paragraph(paragraph) -> bool:
    """判断段落是否为项目符号/编号列表，尽量保留其原有列表样式和缩进。"""
    style_name = ""
    if paragraph.style is not None:
        style_name = (paragraph.style.name or "").lower()

    if style_name.startswith("list"):
        return True

    p_pr = paragraph._element.pPr
    return p_pr is not None and p_pr.numPr is not None


def _clear_paragraph_numbering(paragraph):
    """移除段落原有的编号定义，避免旧模板列表样式残留。"""
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return

    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def iter_table_paragraphs(tables):
    """递归遍历所有表格单元格内的段落。"""
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                yield from iter_table_paragraphs(cell.tables)


def iter_all_tables(tables):
    """递归遍历文档中的所有表格（含嵌套表格）。"""
    for table in tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from iter_all_tables(cell.tables)


def _has_drawing(paragraph) -> bool:
    """判断段落中是否包含图片等 drawing 元素。"""
    return _element_contains_any_tag(paragraph._element, DRAWING_XML_TAGS)


def _has_equation_content(paragraph) -> bool:
    """判断段落中是否包含 Word 公式或嵌入式公式对象。"""
    return _element_contains_any_tag(paragraph._element, EQUATION_XML_TAGS)


def _element_contains_any_tag(element, tags: frozenset[str]) -> bool:
    """判断 XML 子树中是否包含任一目标标签。"""
    return any(getattr(node, "tag", None) in tags for node in element.iter())


def _scan_paragraph_content(paragraph) -> tuple[bool, bool]:
    """单次扫描段落 XML，同时探测图片/对象与公式内容。"""
    has_drawing = False
    has_equation = False

    for node in paragraph._element.iter():
        tag = getattr(node, "tag", None)
        if tag in DRAWING_XML_TAGS:
            has_drawing = True
        if tag in EQUATION_XML_TAGS:
            has_equation = True
        if has_drawing and has_equation:
            break

    return has_drawing, has_equation


def _ensure_xml_child(parent, tag_name: str, *, prepend: bool = False):
    """确保底层 XML 节点存在，便于对 DOCX 压缩包中的部件做后处理。"""
    child = parent.find(qn(tag_name))
    if child is not None:
        return child

    child = OxmlElement(tag_name)
    if prepend:
        parent.insert(0, child)
    else:
        parent.append(child)
    return child


def _set_xml_attribute(element, attr_name: str, value: str) -> bool:
    """仅在属性值变化时写入，便于统计 XML 是否被修改。"""
    attr = qn(attr_name)
    if element.get(attr) == value:
        return False

    element.set(attr, value)
    return True


def _format_footnote_run_xml(run) -> bool:
    """统一脚注 run 的中西文字体和字号。"""
    changed = False
    r_pr = run.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        run.insert(0, r_pr)
        changed = True

    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
        changed = True

    changed |= _set_xml_attribute(r_fonts, "w:eastAsia", "宋体")
    changed |= _set_xml_attribute(r_fonts, "w:ascii", "Times New Roman")
    changed |= _set_xml_attribute(r_fonts, "w:hAnsi", "Times New Roman")
    changed |= _set_xml_attribute(r_fonts, "w:cs", "Times New Roman")

    size_val = str(int(FOOTNOTE_FONT_SIZE_PT * 2))
    size = _ensure_xml_child(r_pr, "w:sz")
    size_cs = _ensure_xml_child(r_pr, "w:szCs")
    changed |= _set_xml_attribute(size, "w:val", size_val)
    changed |= _set_xml_attribute(size_cs, "w:val", size_val)

    if run.find(qn("w:footnoteRef")) is not None or run.find(qn("w:footnoteReference")) is not None:
        r_style = _ensure_xml_child(r_pr, "w:rStyle", prepend=True)
        vert_align = _ensure_xml_child(r_pr, "w:vertAlign")
        changed |= _set_xml_attribute(r_style, "w:val", "FootnoteReference")
        changed |= _set_xml_attribute(vert_align, "w:val", "superscript")

    return changed


def _format_footnote_paragraph_xml(paragraph) -> bool:
    """统一脚注段落的行距、对齐和缩进。"""
    changed = False
    p_pr = paragraph.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph.insert(0, p_pr)
        changed = True

    spacing = _ensure_xml_child(p_pr, "w:spacing")
    changed |= _set_xml_attribute(spacing, "w:before", "0")
    changed |= _set_xml_attribute(spacing, "w:after", "0")
    changed |= _set_xml_attribute(spacing, "w:line", "240")
    changed |= _set_xml_attribute(spacing, "w:lineRule", "auto")

    ind = _ensure_xml_child(p_pr, "w:ind")
    changed |= _set_xml_attribute(ind, "w:left", "0")
    changed |= _set_xml_attribute(ind, "w:right", "0")
    changed |= _set_xml_attribute(ind, "w:firstLine", "0")
    if ind.get(qn("w:hanging")) is not None:
        del ind.attrib[qn("w:hanging")]
        changed = True

    jc = _ensure_xml_child(p_pr, "w:jc")
    changed |= _set_xml_attribute(jc, "w:val", "left")

    widow_control = _ensure_xml_child(p_pr, "w:widowControl")
    changed |= _set_xml_attribute(widow_control, "w:val", "true")

    return changed


def _rewrite_docx_part(docx_path: str | Path, part_name: str, transform) -> int:
    """
    重写 docx 中指定部件。

    transform 回调返回 `(new_bytes, count, changed)`。
    """
    docx_file = Path(docx_path)
    if not docx_file.exists():
        return 0

    with ZipFile(docx_file, "r") as source:
        if part_name not in source.namelist():
            return 0

        entries = source.infolist()
        payloads = {entry.filename: source.read(entry.filename) for entry in entries}

    new_bytes, count, changed = transform(payloads[part_name])
    if not changed:
        return count

    payloads[part_name] = new_bytes

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        temp_path = Path(handle.name)

    try:
        with ZipFile(temp_path, "w") as target:
            for entry in entries:
                target.writestr(entry, payloads[entry.filename])
        temp_path.replace(docx_file)
    finally:
        temp_path.unlink(missing_ok=True)

    return count


def format_docx_footnotes(docx_path: str | Path) -> int:
    """
    统一脚注正文的字体与字号。

    `python-docx` 目前缺少稳定的脚注公开 API，因此这里在文档保存后
    直接修正 `word/footnotes.xml` 中的 run 属性，尽量以最小改动覆盖
    真实论文里常见的“脚注字号/字体不统一”问题。
    """
    def transform(xml_bytes: bytes):
        root = etree.fromstring(xml_bytes)
        formatted_count = 0
        changed = False

        for footnote in root.findall(qn("w:footnote")):
            footnote_type = footnote.get(qn("w:type"))
            footnote_id = footnote.get(qn("w:id"))
            if footnote_type in FOOTNOTE_SKIP_TYPES:
                continue

            if footnote_id is not None:
                try:
                    if int(footnote_id) < 0:
                        continue
                except ValueError:
                    pass

            footnote_changed = False
            for paragraph in footnote.findall("./" + qn("w:p")):
                footnote_changed |= _format_footnote_paragraph_xml(paragraph)
            for run in footnote.findall(".//" + qn("w:r")):
                footnote_changed |= _format_footnote_run_xml(run)

            if footnote_changed:
                formatted_count += 1
                changed = True

        xml_output = etree.tostring(
            root,
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )
        return xml_output, formatted_count, changed

    try:
        return _rewrite_docx_part(docx_path, FOOTNOTE_XML_PATH, transform)
    except Exception as exc:
        logger.warning(f"统一脚注格式时出现警告：{exc}")
        return 0


def _set_paragraph_format(
    paragraph,
    alignment=None,
    first_line_indent=None,
    space_before=None,
    space_after=None,
    line_spacing=None,
    line_spacing_rule=None,
):
    """
    设置段落格式属性。

    Args:
        paragraph:          docx Paragraph 对象
        alignment:          对齐方式（WD_ALIGN_PARAGRAPH 枚举值）
        first_line_indent:  首行缩进（Cm/Pt 等 docx.shared 对象）
        space_before:       段前间距
        space_after:        段后间距
        line_spacing:       行距数值
        line_spacing_rule:  行距规则（WD_LINE_SPACING 枚举值）
    """
    pf = paragraph.paragraph_format

    if alignment is not None:
        pf.alignment = alignment

    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    else:
        # 显式清除首行缩进（避免继承模板样式）
        pf.first_line_indent = None

    if space_before is not None:
        pf.space_before = space_before

    if space_after is not None:
        pf.space_after = space_after

    if line_spacing is not None:
        pf.line_spacing = line_spacing

    if line_spacing_rule is not None:
        pf.line_spacing_rule = line_spacing_rule


def _set_paragraph_outline_level(paragraph, level: int | None):
    """为段落设置目录级别，便于 Word TOC 字段收录。"""
    p_pr = paragraph._element.get_or_add_pPr()
    outline_level = p_pr.find(qn("w:outlineLvl"))

    if level is None:
        if outline_level is not None:
            p_pr.remove(outline_level)
        return

    if outline_level is None:
        outline_level = OxmlElement("w:outlineLvl")
        p_pr.append(outline_level)

    outline_level.set(qn("w:val"), str(level))


def _set_paragraph_on_off_flag(paragraph, flag: str, enabled: bool):
    """设置段落级 on/off 标志，如 keepNext、keepLines。"""
    p_pr = paragraph._element.get_or_add_pPr()
    flag_element = p_pr.find(qn(f"w:{flag}"))

    if not enabled:
        if flag_element is not None:
            p_pr.remove(flag_element)
        return

    if flag_element is None:
        flag_element = OxmlElement(f"w:{flag}")
        p_pr.append(flag_element)

    flag_element.set(qn("w:val"), "true")


def _set_paragraph_pagination_flags(
    paragraph,
    *,
    keep_next: bool | None = None,
    keep_lines: bool | None = None,
    widow_control: bool | None = None,
):
    """统一设置段落分页相关控制项。"""
    if keep_next is not None:
        _set_paragraph_on_off_flag(paragraph, "keepNext", keep_next)

    if keep_lines is not None:
        _set_paragraph_on_off_flag(paragraph, "keepLines", keep_lines)

    if widow_control is not None:
        _set_paragraph_on_off_flag(paragraph, "widowControl", widow_control)


def _clear_paragraph_style(paragraph, preserve_list_style: bool = False):
    """
    清除段落的已有样式设置，防止模板样式干扰排版。
    将段落样式重置为 Normal。
    """
    if preserve_list_style and _is_list_paragraph(paragraph):
        return

    try:
        paragraph.style = "Normal"
    except KeyError:
        # 某些第三方模板/导出的 docx 可能缺少内置 Normal 样式；
        # 这时改为移除显式段落样式，避免整个排版流程直接失败。
        paragraph.style = None
    _clear_paragraph_numbering(paragraph)


def _get_primary_paragraph(container):
    """获取页眉/页脚中的主段落，并清理多余段落。"""
    paragraphs = list(container.paragraphs)
    paragraph = paragraphs[0] if paragraphs else container.add_paragraph()

    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    _remove_all_runs(paragraph)
    return paragraph


def _get_primary_cell_paragraph(cell):
    """获取单元格中的主段落，并删除多余的默认空段落。"""
    paragraphs = list(cell.paragraphs)
    paragraph = paragraphs[0] if paragraphs else cell.add_paragraph()

    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    _remove_all_runs(paragraph)
    return paragraph


def _get_next_abstract_num_id(numbering_root) -> int:
    """返回 numbering.xml 中下一个可用的 abstractNumId。"""
    abstract_ids = []
    for abstract_num in numbering_root.findall("./" + qn("w:abstractNum")):
        raw_id = abstract_num.get(qn("w:abstractNumId"))
        if raw_id is not None:
            abstract_ids.append(int(raw_id))

    return max(abstract_ids, default=-1) + 1


def _get_or_create_heading_numbering_abstract_id(doc) -> int:
    """获取学术论文标题专用的多级编号定义。"""
    cached_id = getattr(doc, "_academic_heading_abstract_num_id", None)
    if cached_id is not None:
        return cached_id

    numbering_root = doc.part.numbering_part.numbering_definitions._numbering

    for abstract_num in numbering_root.findall("./" + qn("w:abstractNum")):
        nsid = abstract_num.find(qn("w:nsid"))
        if nsid is not None and nsid.get(qn("w:val")) == HEADING_NUMBERING_NSID:
            abstract_num_id = int(abstract_num.get(qn("w:abstractNumId")))
            setattr(doc, "_academic_heading_abstract_num_id", abstract_num_id)
            return abstract_num_id

    abstract_num_id = _get_next_abstract_num_id(numbering_root)
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), HEADING_NUMBERING_NSID)
    abstract_num.append(nsid)

    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level_type)

    template = OxmlElement("w:tmpl")
    template.set(qn("w:val"), HEADING_NUMBERING_TEMPLATE)
    abstract_num.append(template)

    for ilvl, level_text in enumerate(HEADING_NUMBERING_LEVEL_TEXTS):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        lvl.append(num_fmt)

        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "space")
        lvl.append(suffix)

        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), level_text)
        lvl.append(lvl_text)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        abstract_num.append(lvl)

    children = list(numbering_root)
    insert_at = next((index for index, child in enumerate(children) if child.tag == qn("w:num")), len(children))
    numbering_root.insert(insert_at, abstract_num)
    setattr(doc, "_academic_heading_abstract_num_id", abstract_num_id)
    return abstract_num_id


def _create_heading_numbering_instance(doc, numbering_parts: tuple[int, ...]) -> int | None:
    """为当前标题创建一个 concrete numbering 实例，并保留原始章节号。"""
    if not numbering_parts:
        return None

    abstract_num_id = _get_or_create_heading_numbering_abstract_id(doc)
    numbering_root = doc.part.numbering_part.numbering_definitions._numbering
    num = numbering_root.add_num(abstract_num_id)

    for ilvl, start_value in enumerate(numbering_parts):
        lvl_override = num.add_lvlOverride(ilvl)
        lvl_override.add_startOverride(start_value)

    return int(num.get(qn("w:numId")))


def _apply_paragraph_numbering(paragraph, num_id: int, ilvl: int):
    """将 concrete numbering 绑定到段落。"""
    p_pr = paragraph._element.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))

    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    else:
        for child in list(num_pr):
            num_pr.remove(child)

    ilvl_element = OxmlElement("w:ilvl")
    ilvl_element.set(qn("w:val"), str(ilvl))
    num_pr.append(ilvl_element)

    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_element)


def apply_native_heading_numbering(doc, paragraph, para_type: str, numbering_parts: tuple[int, ...]):
    """把识别出的章节号写成 Word 原生多级编号。"""
    level = HEADING_LEVEL_BY_TYPE.get(para_type)
    if level is None:
        return

    num_id = _create_heading_numbering_instance(doc, numbering_parts)
    if num_id is None:
        return

    _apply_paragraph_numbering(paragraph, num_id=num_id, ilvl=level)


def _set_xml_borders(border_container, border_map: dict[str, dict[str, str]]):
    """
    使用底层 XML 设置表格/单元格边框。

    python-docx 没有提供“只保留某一条边框”的高级接口，
    因此需要直接写 WordprocessingML 的边框节点属性。
    """
    for edge, attrs in border_map.items():
        edge_tag = qn(f"w:{edge}")
        border = border_container.find(edge_tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            border_container.append(border)

        for key, value in attrs.items():
            border.set(qn(f"w:{key}"), str(value))


def _hidden_border_attrs() -> dict[str, str]:
    """
    返回隐藏边框使用的属性。

    `none` 在 WPS 对表格/段落边框的兼容性通常比 `nil` 更稳定，
    更适合“隐藏大部分边框、只显示个别一条线”的场景。
    """
    return {
        "val": "none",
        "sz": "0",
        "space": "0",
        "color": "auto",
    }


def _hide_table_borders(table):
    """
    在表格级别隐藏所有边框。

    先把整张表的外框和内部横竖线全部置为 nil，
    后面再对“值列”的单元格单独开启 bottom 边框，
    这样就能得到类似“填写横线”的封面效果。
    """
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    hidden = _hidden_border_attrs()
    _set_xml_borders(
        tbl_borders,
        {
            "top": hidden,
            "left": hidden,
            "bottom": hidden,
            "right": hidden,
            "insideH": hidden,
            "insideV": hidden,
        },
    )


def _hide_cell_borders(cell):
    """显式隐藏单元格四周边框，避免模板样式残留。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    hidden = _hidden_border_attrs()
    _set_xml_borders(
        tc_borders,
        {
            "top": hidden,
            "left": hidden,
            "right": hidden,
            "bottom": hidden,
        },
    )


def _set_cell_only_bottom_border(cell, color: str = "000000", size: str = "8"):
    """
    只保留单元格的下边框，用来模拟封面信息栏的填写横线。

    实现步骤：
    1. 先把 top/left/right/bottom 全部清空为 nil
    2. 再把 bottom 单独改成 single

    这样可以确保无论模板原本是否带边框，最终都只有底部这一条线可见。
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    hidden = _hidden_border_attrs()
    _set_xml_borders(
        tc_borders,
        {
            "top": hidden,
            "left": hidden,
            "right": hidden,
            "bottom": hidden,
        },
    )
    _set_xml_borders(
        tc_borders,
        {
            "bottom": {
                "val": "single",
                "sz": size,
                "space": "0",
                "color": color,
            }
        },
    )


def _set_explicit_cell_borders(cell, top: bool = False, bottom: bool = False, color: str = "000000", size: str = "10"):
    """
    通过直接覆写单元格级别的四面边框，强行阻断并覆盖所有来自表格级别的样式继承。
    对于不需要显示的边，设置为 val="none"。
    这能保证 100% 出现需要的三线表线条，并消除表格自带格线。
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is not None:
        tc_pr.remove(tc_borders)
        
    tc_borders = OxmlElement("w:tcBorders")
    tc_pr.append(tc_borders)

    hidden = _hidden_border_attrs()
    visible = {
        "val": "single",
        "sz": size,
        "space": "0",
        "color": color,
    }
    
    _set_xml_borders(
        tc_borders,
        {
            "top": visible if top else hidden,
            "bottom": visible if bottom else hidden,
            "left": hidden,
            "right": hidden,
        },
    )

def _remove_table_borders(table):
    """移除表格级的边框定义，以防与单元格边框交织冲突"""
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is not None:
        tbl_pr.remove(tbl_borders)


def _set_row_repeat_as_header(row, enabled: bool = True):
    """将表格首行标记为跨页重复表头。"""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))

    if not enabled:
        if tbl_header is not None:
            tr_pr.remove(tbl_header)
        return

    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)

    tbl_header.set(qn("w:val"), "true")


def _set_row_cant_split(row, enabled: bool = True):
    """禁止表格行在分页处被拆开，提升长表阅读连续性。"""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))

    if not enabled:
        if cant_split is not None:
            tr_pr.remove(cant_split)
        return

    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

    cant_split.set(qn("w:val"), "true")


def format_three_line_table(table):
    """将普通表格处理为学术论文常见的三线表边框。"""
    rows = list(table.rows)
    if not rows:
        return

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _remove_table_borders(table)
    # 取消底层的表格样式以防底层隐藏线逻辑有奇怪的行为
    if table.style:
        try:
            table.style = "Normal Table"
        except Exception:
            pass

    last_row_index = len(rows) - 1

    for row_index, row in enumerate(rows):
        is_header = (row_index == 0)
        is_last = (row_index == last_row_index)
        _set_row_repeat_as_header(row, enabled=is_header)
        _set_row_cant_split(row, enabled=True)
        
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            for paragraph in cell.paragraphs:
                _set_paragraph_format(
                    paragraph,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_indent=Pt(0),
                    line_spacing=1.5,
                    line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
                )
                if is_header:
                    _apply_run_fonts(
                        paragraph,
                        cn_font="宋体",
                        en_font="Times New Roman",
                        size_pt=12,
                        bold=True,
                    )

            # 第一行：画顶线和底线（粗一点或者普通的单线条）
            # 最后一行：画底线
            # 其他行：全部 none
            _set_explicit_cell_borders(
                cell,
                top=is_header,
                bottom=is_header or is_last,
                size="10" if is_header or is_last else "0",
            )


def _set_paragraph_only_bottom_border(paragraph, color: str = "000000", size: str = "10"):
    """
    给段落仅保留一条下边框。

    WPS 对“单元格边框 + 整表隐藏边框”的组合并不总是稳定，
    所以这里在值列段落上再补一条段落下边框，作为更稳的可视化横线。
    """
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    hidden = _hidden_border_attrs()
    _set_xml_borders(
        p_bdr,
        {
            "top": hidden,
            "left": hidden,
            "right": hidden,
            "bottom": {
                "val": "single",
                "sz": size,
                "space": "1",
                "color": color,
            },
        },
    )


def _prepend_block_elements(doc, elements):
    """
    将一组已创建好的段落/表格 XML 元素移动到文档最前方。

    python-docx 没有“在开头插入 block”的公开 API，
    所以这里采用“先追加、再搬到开头”的方式实现。
    """
    body = doc._body._element

    for insert_index, element in enumerate(elements):
        body.remove(element)
        body.insert(insert_index, element)


def _insert_block_elements_after(paragraph, elements):
    """将一组 block 元素插入到指定段落之后。"""
    parent = paragraph._element.getparent()
    insert_index = parent.index(paragraph._element) + 1

    for offset, element in enumerate(elements):
        element.getparent().remove(element)
        parent.insert(insert_index + offset, element)


def _resolve_cover_asset_path(explicit_path, asset_key: str) -> Path | None:
    """
    解析封面素材路径。

    优先使用显式传入的绝对/相对路径；如果没传，则自动在 `static/` 下
    查找约定好的默认文件名，减少调用方配置成本。
    """
    if explicit_path:
        candidate = Path(str(explicit_path)).expanduser()
        if candidate.exists():
            return candidate

    static_dir = Path(__file__).resolve().parent / "static"
    for filename in COVER_IMAGE_CANDIDATES.get(asset_key, ()):
        candidate = static_dir / filename
        if candidate.exists():
            return candidate

    return None


def _append_page_number_field(paragraph):
    """在页脚插入 Word 自动页码字段。"""
    size_pt = 10.5

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_begin = paragraph.add_run()
    _set_run_font(run_begin, cn_font="宋体", en_font="Times New Roman", size_pt=size_pt)
    run_begin._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = "PAGE"
    run_instr = paragraph.add_run()
    _set_run_font(run_instr, cn_font="宋体", en_font="Times New Roman", size_pt=size_pt)
    run_instr._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run_separate = paragraph.add_run()
    _set_run_font(run_separate, cn_font="宋体", en_font="Times New Roman", size_pt=size_pt)
    run_separate._r.append(fld_char_separate)

    run_text = paragraph.add_run("1")
    _set_run_font(run_text, cn_font="宋体", en_font="Times New Roman", size_pt=size_pt)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end = paragraph.add_run()
    _set_run_font(run_end, cn_font="宋体", en_font="Times New Roman", size_pt=size_pt)
    run_end._r.append(fld_char_end)


def _paragraph_contains_page_break(paragraph) -> bool:
    """判断段落里是否已经包含显式分页符。"""
    return 'w:type="page"' in paragraph._element.xml


def _append_forced_page_break(doc):
    """在文档末尾追加一个显式分页符。"""
    page_break = doc.add_paragraph()
    _clear_paragraph_style(page_break)
    _set_paragraph_format(
        page_break,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    page_break.add_run().add_break(WD_BREAK.PAGE)
    return page_break


def _prepend_forced_page_break(doc):
    """在文档开头插入一个独立的显式分页段落。"""
    page_break = _append_forced_page_break(doc)
    _prepend_block_elements(doc, [page_break._element])
    return page_break


def ensure_document_ends_with_page_break(doc):
    """确保当前文档末尾带有分页符，便于正文从下一页开始。"""
    if doc.paragraphs and _paragraph_contains_page_break(doc.paragraphs[-1]):
        return

    _append_forced_page_break(doc)


def ensure_document_starts_with_page_break(doc):
    """确保当前文档开头带有分页符，便于内容从新页开始。"""
    if doc.paragraphs and _paragraph_contains_page_break(doc.paragraphs[0]):
        return

    _prepend_forced_page_break(doc)


def get_max_printable_width(doc) -> int:
    """返回文档各节中可打印区域的最小宽度（EMU）。"""
    widths = []
    for section in doc.sections:
        printable_width = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
        if printable_width > 0:
            widths.append(printable_width)

    return min(widths) if widths else 0


def constrain_inline_images(doc, progress_callback=None) -> int:
    """
    检测文档中的内嵌图片，若超出页面可打印宽度则按比例缩小。

    这里只处理 inline_shapes：
    - 对当前项目已经覆盖的“插入式图片”最有效
    - 与 python-docx 的能力边界一致，兼容性相对稳定
    """
    max_width = get_max_printable_width(doc)
    if max_width <= 0:
        return 0

    inline_shapes = list(doc.inline_shapes)
    resized_count = 0
    total_images = len(inline_shapes)

    for index, shape in enumerate(inline_shapes, start=1):
        emit_progress(
            progress_callback,
            3,
            f"正在检查第 {index}/{total_images} 张图片尺寸",
        )
        original_width = int(shape.width)
        original_height = int(shape.height)
        if original_width <= 0 or original_height <= 0 or original_width <= max_width:
            continue

        scale = max_width / original_width
        shape.width = max_width
        shape.height = max(1, int(round(original_height * scale)))
        resized_count += 1
        emit_progress(
            progress_callback,
            3,
            f"已缩小第 {index}/{total_images} 张图片",
            "嵌入型图片宽度已限制到页面可打印区域内",
        )

    # 兼容各种由于直接粘贴导致的非标准图片（如浮动图形 wp:anchor 或旧版 v:shape）
    for anchor in doc._element.findall(".//" + qn("wp:anchor")):
        extent = anchor.find(".//" + qn("wp:extent"))
        if extent is not None:
            cx = int(extent.get("cx", 0))
            cy = int(extent.get("cy", 0))
            if cx > max_width:
                scale = max_width / cx
                new_cx = max_width
                new_cy = max(1, int(round(cy * scale)))
                extent.set("cx", str(new_cx))
                extent.set("cy", str(new_cy))
                for a_ext in anchor.findall(".//" + qn("a:ext")):
                    a_ext.set("cx", str(new_cx))
                    a_ext.set("cy", str(new_cy))
                resized_count += 1

    # 兼容直接从部分网页/Excel 粘贴带来的 VML 图形 (v:shape)
    for v_shape in doc._element.findall(".//{urn:schemas-microsoft-com:vml}shape"):
        style_str = v_shape.get("style", "")
        if "width:" in style_str and "height:" in style_str:
            import re
            w_match = re.search(r"width:([0-9.]+)([a-zA-Z]+)", style_str)
            h_match = re.search(r"height:([0-9.]+)([a-zA-Z]+)", style_str)
            if w_match and h_match:
                w_val = float(w_match.group(1))
                w_unit = w_match.group(2)
                h_val = float(h_match.group(1))
                h_unit = h_match.group(2)
                unit_to_emu = {"pt": 12700, "in": 914400, "cm": 360000, "mm": 36000, "px": 9525}
                w_emu = int(w_val * unit_to_emu.get(w_unit, 12700))
                if w_emu > max_width:
                    scale = max_width / w_emu
                    new_style = re.sub(r"width:[0-9.]+[a-zA-Z]+", f"width:{w_val * scale:.2f}{w_unit}", style_str)
                    new_style = re.sub(r"height:[0-9.]+[a-zA-Z]+", f"height:{h_val * scale:.2f}{h_unit}", new_style)
                    v_shape.set("style", new_style)
                    resized_count += 1

    return resized_count


def _append_toc_field(paragraph):
    """插入 Word 目录域，打开文档更新域后即可生成目录。"""
    size_pt = 12

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_begin = paragraph.add_run()
    _set_run_font(run_begin, cn_font="宋体", en_font="Times New Roman", size_pt=size_pt)
    run_begin._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    run_instr = paragraph.add_run()
    _set_run_font(run_instr, cn_font="宋体", en_font="Times New Roman", size_pt=size_pt)
    run_instr._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run_separate = paragraph.add_run()
    _set_run_font(run_separate, cn_font="宋体", en_font="Times New Roman", size_pt=size_pt)
    run_separate._r.append(fld_char_separate)

    placeholder = paragraph.add_run("目录将在打开文档后自动生成，可右键更新目录立即刷新。")
    _set_run_font(placeholder, cn_font="宋体", en_font="Times New Roman", size_pt=size_pt)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end = paragraph.add_run()
    _set_run_font(run_end, cn_font="宋体", en_font="Times New Roman", size_pt=size_pt)
    run_end._r.append(fld_char_end)


def _enable_field_updates_on_open(doc):
    """提示 Word/WPS 在打开文档时更新域代码。"""
    settings_element = doc.settings.element
    update_fields = settings_element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings_element.append(update_fields)

    update_fields.set(qn("w:val"), "true")


def apply_document_layout(doc, header_text: str) -> dict:
    """统一设置纸张、页边距、页眉和页码。"""
    raw_header = (header_text or "").strip()
    if not raw_header:
        header_value = DEFAULT_HEADER_TEXT
    elif len(raw_header) <= RUNNING_HEADER_MAX_LENGTH:
        header_value = raw_header
    else:
        header_value = f"{raw_header[:RUNNING_HEADER_MAX_LENGTH - 3].rstrip()}..."

    for section in doc.sections:
        section.page_width = Cm(PAGE_LAYOUT["page_width_cm"])
        section.page_height = Cm(PAGE_LAYOUT["page_height_cm"])
        section.top_margin = Cm(PAGE_LAYOUT["margins_cm"]["top"])
        section.bottom_margin = Cm(PAGE_LAYOUT["margins_cm"]["bottom"])
        section.left_margin = Cm(PAGE_LAYOUT["margins_cm"]["left"])
        section.right_margin = Cm(PAGE_LAYOUT["margins_cm"]["right"])
        section.header_distance = Cm(PAGE_LAYOUT["header_distance_cm"])
        section.footer_distance = Cm(PAGE_LAYOUT["footer_distance_cm"])
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        header_paragraph = _get_primary_paragraph(section.header)
        _clear_paragraph_style(header_paragraph)
        _set_paragraph_format(
            header_paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Pt(0),
            line_spacing=1.0,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        if header_value:
            header_run = header_paragraph.add_run(header_value)
            _set_run_font(header_run, cn_font="宋体", en_font="Times New Roman", size_pt=10.5)

        footer_paragraph = _get_primary_paragraph(section.footer)
        _clear_paragraph_style(footer_paragraph)
        _set_paragraph_format(
            footer_paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Pt(0),
            line_spacing=1.0,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        _append_page_number_field(footer_paragraph)

    return {
        "page_size": PAGE_LAYOUT["page_size"],
        "page_width_cm": PAGE_LAYOUT["page_width_cm"],
        "page_height_cm": PAGE_LAYOUT["page_height_cm"],
        "margins_cm": PAGE_LAYOUT["margins_cm"].copy(),
        "header_distance_cm": PAGE_LAYOUT["header_distance_cm"],
        "footer_distance_cm": PAGE_LAYOUT["footer_distance_cm"],
        "header_text": header_value,
        "page_number_position": "footer_center",
    }


def generate_cover_page(doc, info_dict):
    """
    在文档最前方插入一页标准课程论文封面，并在封面末尾补分页符。

    设计为“正文排版完成后再调用”，这样封面标题不会被正文识别逻辑再次改写。
    如果缺少 title，则视为没有封面信息，直接返回 False。
    """
    if doc is None:
        raise ValueError("doc 不能为空")

    info_dict = info_dict or {}
    title_text = normalize_text_for_matching(str(info_dict.get("title", "")))
    cover_title = normalize_text_for_matching(
        str(info_dict.get("cover_title") or info_dict.get("course_title") or title_text)
    )
    if not title_text:
        return False

    new_blocks = []

    school_name = normalize_text_for_matching(str(info_dict.get("school_name", "浙江工商大学")))
    logo_file = _resolve_cover_asset_path(info_dict.get("logo_path"), "logo")
    school_name_image_file = _resolve_cover_asset_path(info_dict.get("school_name_image_path"), "school_name")

    # ---------- 1. 顶部校徽 ----------
    logo_paragraph = doc.add_paragraph()
    _clear_paragraph_style(logo_paragraph)
    _set_paragraph_format(
        logo_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        space_before=Pt(COVER_LAYOUT["logo_space_before_pt"]),
        space_after=Pt(COVER_LAYOUT["logo_space_after_pt"]),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    if logo_file is not None:
        logo_paragraph.add_run().add_picture(str(logo_file), width=Cm(COVER_LAYOUT["logo_width_cm"]))
    new_blocks.append(logo_paragraph._element)

    # ---------- 2. 校名 ----------
    school_paragraph = doc.add_paragraph()
    _clear_paragraph_style(school_paragraph)
    _set_paragraph_format(
        school_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        space_after=Pt(COVER_LAYOUT["school_name_space_after_pt"]),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    if school_name_image_file is not None:
        school_paragraph.add_run().add_picture(str(school_name_image_file), width=Cm(COVER_LAYOUT["school_name_width_cm"]))
    else:
        school_run = school_paragraph.add_run(school_name)
        _set_run_font(school_run, cn_font="华文行楷", en_font="Times New Roman", size_pt=28, bold=False)
    new_blocks.append(school_paragraph._element)

    # ---------- 3. 封面大标题 ----------
    title_paragraph = doc.add_paragraph()
    _clear_paragraph_style(title_paragraph)
    _set_paragraph_format(
        title_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        space_after=Pt(COVER_LAYOUT["title_space_after_pt"]),
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    title_run = title_paragraph.add_run(cover_title)
    _set_run_font(
        title_run,
        cn_font="宋体",
        en_font="Times New Roman",
        size_pt=COVER_LAYOUT["title_size_pt"],
        bold=True,
    )
    new_blocks.append(title_paragraph._element)

    info_spacer = doc.add_paragraph()
    _clear_paragraph_style(info_spacer)
    _set_paragraph_format(
        info_spacer,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        space_after=Pt(COVER_LAYOUT["info_spacer_after_pt"]),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    new_blocks.append(info_spacer._element)

    # ---------- 4. 个人信息栏 ----------
    # 绝不使用空格/下划线硬凑对齐，而是借助 5x2 表格完成稳定布局。
    table = doc.add_table(rows=len(COVER_INFO_FIELDS), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _hide_table_borders(table)

    for row_index, (label, key) in enumerate(COVER_INFO_FIELDS):
        label_cell = table.cell(row_index, 0)
        value_cell = table.cell(row_index, 1)

        label_cell.width = Cm(COVER_LAYOUT["label_width_cm"])
        value_cell.width = Cm(COVER_LAYOUT["value_width_cm"])
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        label_paragraph = _get_primary_cell_paragraph(label_cell)
        _clear_paragraph_style(label_paragraph)
        _set_paragraph_format(
            label_paragraph,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            first_line_indent=Pt(0),
            space_before=Pt(6),
            space_after=Pt(6),
            line_spacing=1.25,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        label_run = label_paragraph.add_run(label)
        _set_run_font(
            label_run,
            cn_font="宋体",
            en_font="Times New Roman",
            size_pt=COVER_LAYOUT["info_font_pt"],
            bold=True,
        )

        value_paragraph = _get_primary_cell_paragraph(value_cell)
        _clear_paragraph_style(value_paragraph)
        _set_paragraph_format(
            value_paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Pt(0),
            space_before=Pt(6),
            space_after=Pt(6),
            line_spacing=1.25,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        value_text = normalize_text_for_matching(str(info_dict.get(key, "")))
        value_run = value_paragraph.add_run(value_text)
        # 统一设置中西文字体对，中文保持宋体，数字学号自动显示为 Times New Roman。
        _set_run_font(
            value_run,
            cn_font="宋体",
            en_font="Times New Roman",
            size_pt=COVER_LAYOUT["info_font_pt"],
            bold=False,
        )

        _hide_cell_borders(label_cell)
        _set_cell_only_bottom_border(value_cell, size="10")
        _set_paragraph_only_bottom_border(value_paragraph, size="10")

    new_blocks.append(table._element)

    # ---------- 5. 封面结束后分页 ----------
    page_break = doc.add_paragraph()
    _clear_paragraph_style(page_break)
    _set_paragraph_format(
        page_break,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    page_break.add_run().add_break(WD_BREAK.PAGE)
    new_blocks.append(page_break._element)

    _prepend_block_elements(doc, new_blocks)

    # 使用“首页不同”模式，让封面页不显示页眉和页码，正文从第二页开始承接常规页眉页码。
    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = True

    first_page_header = _get_primary_paragraph(first_section.first_page_header)
    _clear_paragraph_style(first_page_header)
    _set_paragraph_format(
        first_page_header,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )

    first_page_footer = _get_primary_paragraph(first_section.first_page_footer)
    _clear_paragraph_style(first_page_footer)
    _set_paragraph_format(
        first_page_footer,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )

    return True


def prepare_cover_info(cover_info, detected_title: str) -> dict | None:
    """整理自动封面所需信息，并为缺省标题补齐回退值。"""
    if not isinstance(cover_info, dict):
        return None

    resolved = {}
    for key in (
        "title",
        "cover_title",
        "course_title",
        "college",
        "teacher",
        "class_name",
        "student_name",
        "student_id",
        "school_name",
        "logo_path",
        "school_name_image_path",
    ):
        value = cover_info.get(key)
        if value is None:
            continue
        resolved[key] = str(value).strip() if isinstance(value, str) else value

    fallback_title = normalize_text_for_matching(
        str(
            resolved.get("title")
            or detected_title
            or resolved.get("cover_title")
            or resolved.get("course_title")
            or ""
        )
    )
    if not fallback_title:
        return None

    resolved["title"] = fallback_title
    return resolved


def insert_table_of_contents(doc, title_index: int | None, heading_count: int) -> bool:
    """在标题后插入目录标题、TOC 域和分页符。"""
    if title_index is None or heading_count < MIN_TOC_HEADING_COUNT:
        return False

    title_paragraph = doc.paragraphs[title_index]
    new_blocks = []

    toc_heading = doc.add_paragraph()
    _clear_paragraph_style(toc_heading)
    _replace_paragraph_text(toc_heading, "目录")
    _set_paragraph_format(
        toc_heading,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        space_before=Pt(16),
        space_after=Pt(12),
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(toc_heading, None)
    _apply_run_fonts(toc_heading, cn_font="黑体", en_font="Times New Roman", size_pt=16, bold=True)
    new_blocks.append(toc_heading._element)

    toc_paragraph = doc.add_paragraph()
    _clear_paragraph_style(toc_paragraph)
    _set_paragraph_format(
        toc_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(toc_paragraph, None)
    _append_toc_field(toc_paragraph)
    new_blocks.append(toc_paragraph._element)

    page_break = doc.add_paragraph()
    _clear_paragraph_style(page_break)
    _set_paragraph_format(
        page_break,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(page_break, None)
    page_break.add_run().add_break(WD_BREAK.PAGE)
    new_blocks.append(page_break._element)

    _insert_block_elements_after(title_paragraph, new_blocks)
    _enable_field_updates_on_open(doc)
    return True


# ============================================================
# 各类型段落的格式化函数
# ============================================================
def format_body(
    paragraph,
    in_table: bool = False,
    *,
    normalized_text: str | None = None,
    has_equation: bool | None = None,
    has_drawing: bool | None = None,
    english_template_mode: bool = False,
):
    """
    正文格式：
      - 中文字体：宋体
      - 西文字体：Times New Roman
      - 字号：小四（12pt）
      - 首行缩进：2 个中文字符（约 0.74cm × 2 ≈ 对于小四号字约 24pt）
      - 行距：1.5 倍行距
    """
    if normalized_text is None:
        normalized_text = normalize_text_for_matching(paragraph.text)
    _set_paragraph_outline_level(paragraph, None)
    if has_equation is None:
        has_equation = _has_equation_content(paragraph)

    if has_equation and not normalized_text:
        _clear_paragraph_style(paragraph)
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT if in_table else WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Pt(0),
            line_spacing=1.5,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        _set_paragraph_pagination_flags(
            paragraph,
            keep_next=False,
            keep_lines=True,
            widow_control=True,
        )
        _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12)
        return

    if has_drawing is None:
        has_drawing = _has_drawing(paragraph)

    if has_drawing and not normalized_text:
        _set_paragraph_format(
            paragraph,
            alignment=paragraph.paragraph_format.alignment or WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Pt(0),
        )
        _set_paragraph_pagination_flags(
            paragraph,
            keep_next=not in_table,
            keep_lines=True,
            widow_control=True,
        )
        _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12)
        return

    if _is_list_paragraph(paragraph):
        # 列表段落保留原有项目符号/编号样式，仅统一行距和字体。
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.5
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    else:
        is_english_para = is_english_dominant_text(normalized_text)
        _clear_paragraph_style(paragraph)
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT if in_table or is_english_para else WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=Pt(0) if in_table else Pt(
                ENGLISH_TEMPLATE_BODY_FIRST_INDENT_PT if english_template_mode and is_english_para else 24
            ),
            line_spacing=1.5,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )

    _set_paragraph_pagination_flags(paragraph, widow_control=True)
    _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12)


def format_english_front_matter(paragraph, *, bold: bool = False):
    """英文模板封面区：居中、12pt、标题块加粗。"""
    _clear_paragraph_style(paragraph)
    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        space_before=Pt(0),
        space_after=Pt(0),
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_pagination_flags(paragraph, keep_next=False, keep_lines=True, widow_control=True)
    _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=bold)


def format_title(paragraph, text_override: str | None = None):
    """
    论文标题格式：
      - 字体：黑体
      - 字号：小二（18pt）
      - 加粗
      - 居中对齐
      - 段后适当留白，和摘要区隔开
    """
    _clear_paragraph_style(paragraph)
    if text_override is not None:
        _replace_paragraph_text(paragraph, text_override)
    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        space_after=Pt(18),
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)

    _apply_run_fonts(paragraph, cn_font="黑体", en_font="Times New Roman", size_pt=18, bold=True)


def format_heading_l1(
    paragraph,
    text_override: str | None = None,
    outline_level: int | None = 0,
    *,
    english_template_mode: bool = False,
):
    """
    一级标题格式：
      - 字体：黑体
      - 字号：三号（16pt）
      - 加粗
      - 居中对齐
      - 段前段后：各 1 行（对于三号字 16pt，1 行间距 ≈ 16pt）
    """
    heading_text = text_override if text_override is not None else paragraph.text
    is_english_heading = is_english_dominant_text(heading_text)
    _clear_paragraph_style(paragraph)
    if text_override is not None:
        _replace_paragraph_text(paragraph, text_override)
    if english_template_mode and is_english_heading:
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Pt(0),
            space_before=Pt(18),
            space_after=Pt(12),
            line_spacing=1.5,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        _set_paragraph_outline_level(paragraph, outline_level)
        _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)
        _apply_run_fonts(paragraph, cn_font="黑体", en_font="Times New Roman", size_pt=14, bold=True)
        return

    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT if is_english_heading else WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),  # 标题无缩进
        space_before=Pt(16),      # 段前 1 行（三号字高度 16pt）
        space_after=Pt(16),       # 段后 1 行
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(paragraph, outline_level)
    _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)

    _apply_run_fonts(paragraph, cn_font="黑体", en_font="Times New Roman", size_pt=16, bold=True)


def format_heading_l2(paragraph, text_override: str | None = None, *, english_template_mode: bool = False):
    """
    二级标题格式：
      - 字体：黑体
      - 字号：四号（14pt）
      - 加粗
      - 左对齐
      - 段前段后：各 0.5 行（约 7pt）
    """
    _clear_paragraph_style(paragraph)
    if text_override is not None:
        _replace_paragraph_text(paragraph, text_override)
    if english_template_mode and is_english_dominant_text(text_override if text_override is not None else paragraph.text):
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Pt(0),
            space_before=Pt(14),
            space_after=Pt(10),
            line_spacing=1.5,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        _set_paragraph_outline_level(paragraph, 1)
        _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)
        _apply_run_fonts(paragraph, cn_font="黑体", en_font="Times New Roman", size_pt=13, bold=True)
        return

    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),  # 标题无缩进
        space_before=Pt(7),       # 段前 0.5 行（四号字 14pt × 0.5 = 7pt）
        space_after=Pt(7),        # 段后 0.5 行
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(paragraph, 1)
    _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)

    _apply_run_fonts(paragraph, cn_font="黑体", en_font="Times New Roman", size_pt=14, bold=True)


def format_heading_l3(paragraph, text_override: str | None = None, *, english_template_mode: bool = False):
    """
    三级标题格式：
      - 字体：宋体
      - 字号：小四（12pt）
      - 加粗
      - 左对齐
      - 段前段后：各 0.5 行
    """
    _clear_paragraph_style(paragraph)
    if text_override is not None:
        _replace_paragraph_text(paragraph, text_override)
    if english_template_mode and is_english_dominant_text(text_override if text_override is not None else paragraph.text):
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Pt(0),
            space_before=Pt(12),
            space_after=Pt(8),
            line_spacing=1.5,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        _set_paragraph_outline_level(paragraph, 2)
        _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)
        _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=True)
        return

    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),
        space_before=Pt(6),
        space_after=Pt(6),
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(paragraph, 2)
    _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)

    _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=True)


def format_figure_table(paragraph, text_override: str | None = None, *, keep_next: bool = False):
    """
    图表标题格式：
      - 字体：黑体
      - 字号：五号（10.5pt）
      - 居中对齐
      - 取消首行缩进
    """
    _clear_paragraph_style(paragraph)
    if text_override is not None:
        _replace_paragraph_text(paragraph, text_override)
    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),  # 取消首行缩进
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_pagination_flags(paragraph, keep_next=keep_next, keep_lines=True)

    _apply_run_fonts(paragraph, cn_font="黑体", en_font="Times New Roman", size_pt=10.5)


def format_english_split_caption_label(paragraph, text_override: str | None = None):
    """英文模板表题第一行：Table 1 / Figure 1。"""
    _clear_paragraph_style(paragraph)
    if text_override is not None:
        _replace_paragraph_text(paragraph, text_override)
    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),
        space_before=Pt(14),
        space_after=Pt(0),
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)
    _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=True)


def format_english_split_caption_title(paragraph, text_override: str | None = None):
    """英文模板表题第二行：表题正文。"""
    _clear_paragraph_style(paragraph)
    if text_override is not None:
        _replace_paragraph_text(paragraph, text_override)
    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),
        space_before=Pt(0),
        space_after=Pt(6),
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)
    _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=False)


def format_references_heading(
    paragraph,
    text_override: str | None = None,
    *,
    english_template_mode: bool = False,
):
    """
    参考文献标题格式：
      - 居中
      - 宋体加粗
      - 比一级标题更克制，贴近参考样文的紧凑风格
      - 自动补全为“参考文献：”
    """
    _clear_paragraph_style(paragraph)
    heading_text = normalize_text_for_matching(text_override if text_override is not None else paragraph.text)
    if heading_text:
        heading_text = "References" if is_english_dominant_text(heading_text) else "参考文献："
        _replace_paragraph_text(paragraph, heading_text)

    if english_template_mode and is_english_dominant_text(heading_text):
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Pt(0),
            space_before=Pt(18),
            space_after=Pt(12),
            line_spacing=1.5,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        _set_paragraph_outline_level(paragraph, None)
        _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)
        _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=14, bold=True)
        return

    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        space_before=Pt(12),
        space_after=Pt(12),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_pagination_flags(paragraph, keep_next=True, keep_lines=True)

    _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=True)


def _format_labeled_paragraph(
    paragraph,
    label_pattern: re.Pattern,
    *,
    cn_font: str,
    en_font: str,
    size_pt: float,
    alignment,
    first_line_indent,
    label_text_override: str | None = None,
):
    """按“标签 + 正文”结构重建段落，并统一字体与加粗规则。"""
    _clear_paragraph_style(paragraph)
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_format(
        paragraph,
        alignment=alignment,
        first_line_indent=first_line_indent,
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_pagination_flags(paragraph, widow_control=True)

    full_text = paragraph.text
    match = label_pattern.match(full_text)

    if match:
        label_end = match.end()
        label_text = label_text_override if label_text_override is not None else full_text[:label_end]
        body_text = full_text[label_end:]

        _remove_all_runs(paragraph)

        run_label = paragraph.add_run(label_text)
        _set_run_font(run_label, cn_font=cn_font, en_font=en_font, size_pt=size_pt, bold=True)

        if body_text:
            run_body = paragraph.add_run(body_text)
            _set_run_font(run_body, cn_font=cn_font, en_font=en_font, size_pt=size_pt, bold=False)
        return

    _apply_run_fonts(paragraph, cn_font=cn_font, en_font=en_font, size_pt=size_pt)


def format_abstract_or_keywords(paragraph, label_pattern: re.Pattern):
    """
    摘要 / 关键词段落格式：
      - 整体按正文格式（宋体，小四，首行缩进，1.5倍行距）
      - 将标签部分（如 "摘要："、"关键词："）加粗以起强调作用

    实现思路：
      1. 将段落所有 run 的文本拼接
      2. 用正则找到标签的结束位置
      3. 将段落拆分为 "标签 run" 和 "正文 run" 两部分
      4. 标签 run 设为加粗，正文 run 不加粗

    Args:
        paragraph:     docx Paragraph 对象
        label_pattern: 用于匹配标签的正则表达式
    """
    _format_labeled_paragraph(
        paragraph,
        label_pattern,
        cn_font="宋体",
        en_font="Times New Roman",
        size_pt=12,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Pt(24),
    )

def format_english_abstract_heading(paragraph, *, english_template_mode: bool = False):
    """英文摘要标题格式：居中、Times New Roman、12pt、加粗。"""
    _clear_paragraph_style(paragraph)
    _replace_paragraph_text(paragraph, "Abstract")
    if english_template_mode:
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Pt(0),
            space_before=Pt(18),
            space_after=Pt(12),
            line_spacing=1.5,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        _set_paragraph_outline_level(paragraph, None)
        _set_paragraph_pagination_flags(paragraph, widow_control=True)
        _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=14, bold=True)
        return

    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Pt(0),
        space_before=Pt(12),
        space_after=Pt(12),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_pagination_flags(paragraph, widow_control=True)
    _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=True)


def format_english_abstract(
    paragraph,
    label_pattern: re.Pattern | None = None,
    *,
    english_template_mode: bool = False,
):
    """英文摘要正文：Times New Roman、12pt、1.5 倍行距，不额外首行缩进。"""
    if label_pattern is not None:
        _format_labeled_paragraph(
            paragraph,
            label_pattern,
            cn_font="宋体",
            en_font="Times New Roman",
            size_pt=12,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Pt(0),
            label_text_override="Abstract: ",
        )
        return

    _clear_paragraph_style(paragraph)
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),
        line_spacing=1.5,
        line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    )
    _set_paragraph_pagination_flags(paragraph, widow_control=True)
    _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12)


def format_english_keywords(paragraph, *, english_template_mode: bool = False):
    """英文关键词：标签加粗、正文常规，整体左对齐。"""
    if english_template_mode:
        _clear_paragraph_style(paragraph)
        _set_paragraph_outline_level(paragraph, None)
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Pt(ENGLISH_TEMPLATE_BODY_FIRST_INDENT_PT),
            line_spacing=1.5,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        _set_paragraph_pagination_flags(paragraph, widow_control=True)
        _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=False)
        return

    _format_labeled_paragraph(
        paragraph,
        RE_ENGLISH_KEYWORDS,
        cn_font="宋体",
        en_font="Times New Roman",
        size_pt=12,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),
        label_text_override="Keywords: ",
    )


def format_caption_note(paragraph, *, english_template_mode: bool = False):
    """
    图表附注/来源格式：
      - 左对齐
      - 宋体 / Times New Roman，五号（10.5pt）
      - 单倍行距
      - 标签部分加粗，正文常规
    """
    if english_template_mode and is_english_dominant_text(paragraph.text):
        _clear_paragraph_style(paragraph)
        _set_paragraph_outline_level(paragraph, None)
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Pt(0),
            space_before=Pt(4),
            space_after=Pt(12),
            line_spacing=1.5,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        _set_paragraph_pagination_flags(paragraph, widow_control=True)
        _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=10, bold=False)
        return

    _clear_paragraph_style(paragraph)
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(0),
        space_before=Pt(0),
        space_after=Pt(0),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    _set_paragraph_pagination_flags(paragraph, widow_control=True)

    full_text = paragraph.text
    match = RE_CAPTION_NOTE.match(full_text)
    if match:
        label_text = full_text[:match.end()]
        body_text = full_text[match.end():]

        _remove_all_runs(paragraph)

        run_label = paragraph.add_run(label_text)
        _set_run_font(run_label, cn_font="宋体", en_font="Times New Roman", size_pt=10.5, bold=True)

        if body_text:
            run_body = paragraph.add_run(body_text)
            _set_run_font(run_body, cn_font="宋体", en_font="Times New Roman", size_pt=10.5, bold=False)
        return

    _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=10.5)


def format_reference_entry(paragraph, *, english_template_mode: bool = False):
    """
    参考文献条目格式：
      - 左对齐
      - 宋体 / Times New Roman，五号（10pt）
      - 单倍行距
      - 参考样文使用轻微首行缩进，而不是正文式悬挂缩进
    """
    if english_template_mode and is_english_dominant_text(paragraph.text):
        _clear_paragraph_style(paragraph)
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Pt(-ENGLISH_TEMPLATE_REFERENCE_HANGING_PT),
            space_before=Pt(0),
            space_after=Pt(6),
            line_spacing=1.1666666667,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        )
        _set_paragraph_outline_level(paragraph, None)
        _set_paragraph_pagination_flags(paragraph, widow_control=True)
        paragraph.paragraph_format.left_indent = Pt(ENGLISH_TEMPLATE_REFERENCE_HANGING_PT)
        paragraph.paragraph_format.right_indent = Pt(0)
        _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=11)
        return

    _clear_paragraph_style(paragraph)
    _set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Pt(21),
        space_before=Pt(0),
        space_after=Pt(0),
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    _set_paragraph_outline_level(paragraph, None)
    _set_paragraph_pagination_flags(paragraph, widow_control=True)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)

    _apply_run_fonts(paragraph, cn_font="宋体", en_font="Times New Roman", size_pt=10)


def _append_outline_entry(outline: list[dict], para_type: str, text: str):
    """记录识别到的结构化大纲，供前端预览展示。"""
    if not text:
        return

    level = {
        ParagraphType.TITLE: "title",
        ParagraphType.HEADING_L1: "h1",
        ParagraphType.HEADING_L2: "h2",
        ParagraphType.HEADING_L3: "h3",
        ParagraphType.SECTION_HEADING: "section",
        ParagraphType.REFERENCES_HEADING: "references",
    }.get(para_type)

    if level is None:
        return

    outline.append({"level": level, "text": text})


# ============================================================
# 主函数：学术论文排版
# ============================================================
def format_academic_paper(
    input_path: str,
    output_path: str,
    progress_callback=None,
    cover_info=None,
    format_options=None,
) -> dict | bool:
    """
    读取未排版的 .docx 文档，根据学术论文排版规则进行格式重构，保存为新文档。

    Args:
        input_path:  输入文档路径（.docx）
        output_path: 输出文档路径（.docx）

    Returns:
        True 表示排版成功，False 表示处理失败
    """
    # ---------- 1. 文件校验与读取 ----------
    input_file = Path(input_path)

    if not input_file.exists():
        logger.error(f"输入文件不存在：{input_path}")
        return False

    if not input_file.suffix.lower() == ".docx":
        logger.error(f"不支持的文件格式（仅支持 .docx）：{input_file.suffix}")
        return False

    try:
        doc = Document(input_path)
        logger.info(f"成功读取文档：{input_path}（共 {len(doc.paragraphs)} 个段落）")
        emit_progress(
            progress_callback,
            1,
            "文档读取完成，正在解析结构",
            f"共 {len(doc.paragraphs)} 个段落",
        )
    except Exception as e:
        logger.error(f"无法读取文档 {input_path}：{e}")
        return False

    return _process_document(
        doc,
        output_path,
        progress_callback=progress_callback,
        cover_info=cover_info,
        format_options=format_options,
    )

def format_academic_paper_from_text(
    text: str,
    output_path: str,
    progress_callback=None,
    cover_info=None,
    format_options=None,
) -> dict | bool:
    """
    将纯文本内容转换为符合学术论文排版规则的 .docx 文档。

    Args:
        text:        输入的纯文本内容（多行）
        output_path: 输出文档路径（.docx）

    Returns:
        True 表示排版成功，False 表示处理失败
    """
    try:
        doc = Document()
        for line in split_text_to_paragraphs(text):
            doc.add_paragraph(line)
        logger.info(f"成功从文本创建文档（共 {len(doc.paragraphs)} 个段落）")
        emit_progress(
            progress_callback,
            1,
            "文本读取完成，正在生成文档结构",
            f"共 {len(doc.paragraphs)} 个段落",
        )
    except Exception as e:
        logger.error(f"无法从文本创建文档：{e}")
        return False

    return _process_document(
        doc,
        output_path,
        progress_callback=progress_callback,
        cover_info=cover_info,
        format_options=format_options,
    )

def _process_document(doc, output_path: str, progress_callback=None, cover_info=None, format_options=None) -> dict | bool:
    """内部处理逻辑，将 Document 对象排版并保存。"""
    resolved_format_options = resolve_format_options(format_options)

    # ---------- 2. 设置默认文档级字体 ----------
    emit_progress(progress_callback, 1, "正在初始化页面设置与默认样式")
    style = None
    try:
        style = doc.styles["Normal"]
    except KeyError:
        logger.warning("文档缺少 Normal 样式，已跳过默认样式初始化并回退到段落级直接格式。")

    if style is not None:
        try:
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)

            # 设置 Normal 样式的中文字体为宋体
            r_pr = style.element.get_or_add_rPr()
            r_fonts = r_pr.find(qn("w:rFonts"))
            if r_fonts is None:
                r_fonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                r_pr.insert(0, r_fonts)
            r_fonts.set(qn("w:eastAsia"), "宋体")
        except Exception as e:
            logger.warning(f"设置默认样式时出现警告：{e}")

    # ---------- 3. 统计信息 ----------
    stats = {
        ParagraphType.TITLE: 0,
        ParagraphType.HEADING_L1: 0,
        ParagraphType.HEADING_L2: 0,
        ParagraphType.HEADING_L3: 0,
        ParagraphType.FIGURE_CAPTION: 0,
        ParagraphType.TABLE_CAPTION: 0,
        ParagraphType.CAPTION_NOTE: 0,
        ParagraphType.SECTION_HEADING: 0,
        ParagraphType.REFERENCES_HEADING: 0,
        ParagraphType.REFERENCE_ENTRY: 0,
        ParagraphType.ABSTRACT: 0,
        ParagraphType.KEYWORDS: 0,
        ParagraphType.ENGLISH_ABSTRACT_HEADING: 0,
        ParagraphType.ENGLISH_ABSTRACT: 0,
        ParagraphType.ENGLISH_KEYWORDS: 0,
        ParagraphType.BODY: 0,
    }

    paragraphs = list(doc.paragraphs)
    analyses = _build_paragraph_analyses(paragraphs)
    english_template_mode = detect_english_template_mode(paragraphs, analyses)
    english_front_matter_roles = find_english_front_matter_roles(analyses) if english_template_mode else {}
    english_split_caption_roles = find_english_split_caption_roles(analyses) if english_template_mode else {}
    field_para_indices = _find_field_paragraph_indices(paragraphs)
    title_index = None if english_template_mode else find_title_paragraph_index(paragraphs, analyses)
    title_text = ""
    if title_index is not None:
        title_text = analyses[title_index].normalized_text

    page_setup = apply_document_layout(doc, title_text)
    emit_progress(
        progress_callback,
        1,
        "文档结构解析完成",
        f"共 {len(paragraphs)} 个段落，准备识别标题与摘要",
    )
    outline = []
    in_references = False
    in_english_abstract = False
    previous_nonempty_para_type = None
    previous_caption_related = False
    figure_counter = 0
    table_counter = 0
    equation_paragraph_count = 0
    heading_number_state = [0, 0, 0]
    total_paragraphs = len(paragraphs)

    # ---------- 4. 遍历并格式化每个段落 ----------
    emit_progress(progress_callback, 2, "正在识别标题层级与摘要结构")
    for paragraph, analysis in zip(paragraphs, analyses):
        i = analysis.index
        text = analysis.normalized_text

        # 跳过处于 Word 域块内部的段落（如目录、图目录、表目录条目），
        # 避免把目录条目误认为图表标题并重新格式化。
        if i in field_para_indices:
            _log_detected_paragraph("域内段落(跳过)", i, text)
            if text:
                previous_nonempty_para_type = ParagraphType.BODY
            continue

        para_type = ParagraphType.TITLE if i == title_index else analysis.classified_type
        english_front_matter_role = english_front_matter_roles.get(i)
        english_split_caption_role = english_split_caption_roles.get(i)
        inferred_heading = False
        if para_type == ParagraphType.BODY:
            inferred_para_type = analysis.inferred_heading_type
            if inferred_para_type is not None:
                para_type = inferred_para_type
                inferred_heading = True

        if english_split_caption_role == "table_label":
            para_type = ParagraphType.TABLE_CAPTION
        elif english_split_caption_role == "figure_label":
            para_type = ParagraphType.FIGURE_CAPTION

        if analysis.has_equation:
            equation_paragraph_count += 1

        if total_paragraphs and (
            i == 0
            or i == total_paragraphs - 1
            or (i + 1) % max(1, total_paragraphs // 4 or 1) == 0
        ):
            emit_progress(
                progress_callback,
                2,
                f"正在识别第 {i + 1}/{total_paragraphs} 段的结构",
            )

        if para_type == ParagraphType.REFERENCES_HEADING:
            in_references = True
        elif in_references and text:
            if para_type in {
                ParagraphType.TITLE,
                ParagraphType.HEADING_L1,
                ParagraphType.HEADING_L2,
                ParagraphType.HEADING_L3,
                ParagraphType.SECTION_HEADING,
            }:
                in_references = False
            elif para_type == ParagraphType.BODY or analysis.is_reference_entry_candidate:
                para_type = ParagraphType.REFERENCE_ENTRY
            else:
                in_references = False

        if para_type in {ParagraphType.ENGLISH_ABSTRACT_HEADING, ParagraphType.ENGLISH_ABSTRACT}:
            in_english_abstract = True
        elif in_english_abstract and text:
            if para_type == ParagraphType.ENGLISH_KEYWORDS:
                in_english_abstract = False
            elif para_type in {
                ParagraphType.TITLE,
                ParagraphType.HEADING_L1,
                ParagraphType.HEADING_L2,
                ParagraphType.HEADING_L3,
                ParagraphType.FIGURE_CAPTION,
                ParagraphType.TABLE_CAPTION,
                ParagraphType.SECTION_HEADING,
                ParagraphType.REFERENCES_HEADING,
                ParagraphType.ABSTRACT,
                ParagraphType.KEYWORDS,
            }:
                in_english_abstract = False
            else:
                para_type = ParagraphType.ENGLISH_ABSTRACT

        if (
            para_type == ParagraphType.BODY
            and text
            and analysis.is_caption_note_candidate
            and previous_caption_related
        ):
            para_type = ParagraphType.CAPTION_NOTE

        stats[para_type] += 1
        _append_outline_entry(outline, para_type, text)

        current_caption_related = para_type in {
            ParagraphType.FIGURE_CAPTION,
            ParagraphType.TABLE_CAPTION,
            ParagraphType.CAPTION_NOTE,
        } or english_split_caption_role in {"table_title", "figure_title"}

        if english_front_matter_role is not None:
            _log_detected_paragraph("英文模板封面区", i, text)
            format_english_front_matter(paragraph, bold=english_front_matter_role == "title")

        elif english_split_caption_role == "table_title":
            _log_detected_paragraph("英文模板表题正文", i, text)
            format_english_split_caption_title(paragraph, text_override=text)

        elif english_split_caption_role == "figure_title":
            _log_detected_paragraph("英文模板图题正文", i, text)
            format_english_split_caption_title(paragraph, text_override=text)

        elif para_type == ParagraphType.TITLE:
            _log_detected_paragraph("论文标题", i, text)
            format_title(paragraph, text_override=text)

        elif para_type == ParagraphType.HEADING_L1:
            heading_text, explicit_parts = extract_heading_numbering(text, para_type)
            numbering_parts = resolve_heading_numbering_parts(
                para_type,
                explicit_parts,
                heading_number_state,
                allow_auto_numbering=inferred_heading,
            )
            heading_label = "推断一级标题" if inferred_heading else "一级标题"
            _log_detected_paragraph(heading_label, i, text)
            preserve_explicit_text = should_preserve_explicit_heading_text(text, para_type, english_template_mode)
            render_text = heading_text
            if preserve_explicit_text:
                render_text = text if english_template_mode else rebuild_explicit_heading_text(heading_text, explicit_parts)
            format_heading_l1(
                paragraph,
                text_override=render_text,
                english_template_mode=english_template_mode,
            )
            if not preserve_explicit_text:
                apply_native_heading_numbering(doc, paragraph, para_type, numbering_parts)

        elif para_type == ParagraphType.HEADING_L2:
            heading_text, explicit_parts = extract_heading_numbering(text, para_type)
            numbering_parts = resolve_heading_numbering_parts(
                para_type,
                explicit_parts,
                heading_number_state,
                allow_auto_numbering=inferred_heading,
            )
            heading_label = "推断二级标题" if inferred_heading else "二级标题"
            _log_detected_paragraph(heading_label, i, text)
            preserve_explicit_text = should_preserve_explicit_heading_text(text, para_type, english_template_mode)
            render_text = heading_text
            if preserve_explicit_text:
                render_text = text if english_template_mode else rebuild_explicit_heading_text(heading_text, explicit_parts)
            format_heading_l2(
                paragraph,
                text_override=render_text,
                english_template_mode=english_template_mode,
            )
            if not preserve_explicit_text:
                apply_native_heading_numbering(doc, paragraph, para_type, numbering_parts)

        elif para_type == ParagraphType.HEADING_L3:
            heading_text, explicit_parts = extract_heading_numbering(text, para_type)
            numbering_parts = resolve_heading_numbering_parts(
                para_type,
                explicit_parts,
                heading_number_state,
                allow_auto_numbering=inferred_heading,
            )
            heading_label = "推断三级标题" if inferred_heading else "三级标题"
            _log_detected_paragraph(heading_label, i, text)
            preserve_explicit_text = should_preserve_explicit_heading_text(text, para_type, english_template_mode)
            render_text = heading_text
            if preserve_explicit_text:
                render_text = text if english_template_mode else rebuild_explicit_heading_text(heading_text, explicit_parts)
            format_heading_l3(
                paragraph,
                text_override=render_text,
                english_template_mode=english_template_mode,
            )
            if not preserve_explicit_text:
                apply_native_heading_numbering(doc, paragraph, para_type, numbering_parts)

        elif para_type == ParagraphType.FIGURE_CAPTION:
            figure_counter += 1
            caption_match = analysis.caption_match
            caption_text = text if english_split_caption_role == "figure_label" else rebuild_caption_text(
                ParagraphType.FIGURE_CAPTION,
                figure_counter,
                caption_match[1],
            )
            _log_detected_paragraph("图标题", i, caption_text)
            emit_progress(progress_callback, 2, f"识别到第 {figure_counter} 张图片标题")
            if english_split_caption_role == "figure_label":
                format_english_split_caption_label(paragraph, text_override=caption_text)
            else:
                format_figure_table(paragraph, text_override=caption_text, keep_next=False)

        elif para_type == ParagraphType.TABLE_CAPTION:
            table_counter += 1
            caption_match = analysis.caption_match
            caption_text = text if english_split_caption_role == "table_label" else rebuild_caption_text(
                ParagraphType.TABLE_CAPTION,
                table_counter,
                caption_match[1],
            )
            _log_detected_paragraph("表标题", i, caption_text)
            emit_progress(progress_callback, 2, f"识别到第 {table_counter} 张表格标题")
            if english_split_caption_role == "table_label":
                format_english_split_caption_label(paragraph, text_override=caption_text)
            else:
                format_figure_table(paragraph, text_override=caption_text, keep_next=True)

        elif para_type == ParagraphType.SECTION_HEADING:
            _log_detected_paragraph("非编号章节标题", i, text)
            format_heading_l1(
                paragraph,
                text_override=text,
                outline_level=None,
                english_template_mode=english_template_mode,
            )

        elif para_type == ParagraphType.REFERENCES_HEADING:
            _log_detected_paragraph("参考文献标题", i, text)
            emit_progress(progress_callback, 2, "识别到参考文献区域")
            format_references_heading(paragraph, text_override=text, english_template_mode=english_template_mode)

        elif para_type == ParagraphType.REFERENCE_ENTRY:
            _log_detected_paragraph("参考文献条目", i, text)
            format_reference_entry(paragraph, english_template_mode=english_template_mode)

        elif para_type == ParagraphType.ABSTRACT:
            _log_detected_paragraph("摘要段落", i, text)
            format_abstract_or_keywords(paragraph, RE_ABSTRACT)

        elif para_type == ParagraphType.KEYWORDS:
            _log_detected_paragraph("关键词段", i, text)
            format_abstract_or_keywords(paragraph, RE_KEYWORDS)

        elif para_type == ParagraphType.ENGLISH_ABSTRACT_HEADING:
            _log_detected_paragraph("英文摘要标题", i, text)
            format_english_abstract_heading(paragraph, english_template_mode=english_template_mode)

        elif para_type == ParagraphType.ENGLISH_ABSTRACT:
            _log_detected_paragraph("英文摘要", i, text)
            if RE_ENGLISH_ABSTRACT.match(text):
                format_english_abstract(paragraph, RE_ENGLISH_ABSTRACT, english_template_mode=english_template_mode)
            else:
                format_english_abstract(paragraph, english_template_mode=english_template_mode)

        elif para_type == ParagraphType.ENGLISH_KEYWORDS:
            _log_detected_paragraph("英文关键词", i, text)
            format_english_keywords(paragraph, english_template_mode=english_template_mode)

        elif para_type == ParagraphType.CAPTION_NOTE:
            _log_detected_paragraph("图表附注", i, text)
            format_caption_note(paragraph, english_template_mode=english_template_mode)

        else:
            # 正文段落（含空段落）
            format_body(
                paragraph,
                normalized_text=text,
                has_equation=analysis.has_equation,
                has_drawing=analysis.has_drawing,
                english_template_mode=english_template_mode,
            )

        if text:
            previous_nonempty_para_type = para_type
            previous_caption_related = current_caption_related

    # ---------- 5. 处理表格内段落 ----------
    emit_progress(progress_callback, 3, "正在应用排版规则")
    table_paragraph_count = 0
    for paragraph in iter_table_paragraphs(doc.tables):
        table_paragraph_count += 1
        if _has_equation_content(paragraph):
            equation_paragraph_count += 1
        format_body(paragraph, in_table=True)

    all_tables = list(iter_all_tables(doc.tables))
    for table_index, table in enumerate(all_tables, start=1):
        emit_progress(
            progress_callback,
            3,
            f"正在排版第 {table_index}/{len(all_tables)} 张表格",
            "应用三线表边框与表内正文格式",
        )
        format_three_line_table(table)

    heading_count = (
        stats[ParagraphType.HEADING_L1]
        + stats[ParagraphType.HEADING_L2]
        + stats[ParagraphType.HEADING_L3]
    )
    toc_inserted = False
    if resolved_format_options["insert_toc"] and heading_count >= MIN_TOC_HEADING_COUNT:
        emit_progress(progress_callback, 3, "正在插入自动目录字段")
    if resolved_format_options["insert_toc"]:
        toc_inserted = insert_table_of_contents(doc, title_index, heading_count)

    resized_image_count = 0
    if resolved_format_options["resize_images"]:
        resized_image_count = constrain_inline_images(doc, progress_callback=progress_callback)

    cover_generated = False
    formatted_footnote_count = 0
    resolved_cover_info = prepare_cover_info(cover_info, title_text)
    if resolved_cover_info is not None:
        emit_progress(progress_callback, 3, "正在生成课程论文封面")
        cover_generated = generate_cover_page(doc, resolved_cover_info)
        if cover_generated:
            emit_progress(progress_callback, 3, "封面模板已插入文档首页")

    # ---------- 6. 保存输出文档 ----------
    emit_progress(progress_callback, 4, "正在生成输出文档")
    try:
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        if resolved_format_options["format_footnotes"]:
            emit_progress(progress_callback, 4, "正在统一脚注字体与字号")
            formatted_footnote_count = format_docx_footnotes(output_path)
        logger.info(f"排版完成！已保存至：{output_path}")
    except Exception as e:
        logger.error(f"无法保存文档 {output_path}：{e}")
        return False

    # ---------- 7. 输出统计摘要 ----------
    logger.info("=" * 50)
    logger.info("排版统计：")
    logger.info(f"  论文标题：{stats[ParagraphType.TITLE]} 个")
    logger.info(f"  一级标题：{stats[ParagraphType.HEADING_L1]} 个")
    logger.info(f"  二级标题：{stats[ParagraphType.HEADING_L2]} 个")
    logger.info(f"  三级标题：{stats[ParagraphType.HEADING_L3]} 个")
    logger.info(f"  图标题：{stats[ParagraphType.FIGURE_CAPTION]} 个")
    logger.info(f"  表标题：{stats[ParagraphType.TABLE_CAPTION]} 个")
    logger.info(f"  图表附注：{stats[ParagraphType.CAPTION_NOTE]} 个")
    logger.info(f"  非编号章节标题：{stats[ParagraphType.SECTION_HEADING]} 个")
    logger.info(f"  参考文献标题：{stats[ParagraphType.REFERENCES_HEADING]} 个")
    logger.info(f"  参考文献条目：{stats[ParagraphType.REFERENCE_ENTRY]} 条")
    logger.info(f"  摘要段落：{stats[ParagraphType.ABSTRACT]} 个")
    logger.info(f"  关键词段：{stats[ParagraphType.KEYWORDS]} 个")
    logger.info(f"  英文摘要标题：{stats[ParagraphType.ENGLISH_ABSTRACT_HEADING]} 个")
    logger.info(f"  英文摘要段落：{stats[ParagraphType.ENGLISH_ABSTRACT]} 个")
    logger.info(f"  英文关键词：{stats[ParagraphType.ENGLISH_KEYWORDS]} 个")
    logger.info(f"  正文段落：{stats[ParagraphType.BODY]} 个")
    logger.info(f"  表格内段落：{table_paragraph_count} 个")
    logger.info(f"  公式段落：{equation_paragraph_count} 个")
    logger.info(f"  已统一脚注：{formatted_footnote_count} 条")
    logger.info(f"  自动缩放图片：{resized_image_count} 张")
    logger.info(f"  自动封面：{'已生成' if cover_generated else '未生成'}")
    logger.info("=" * 50)

    return {
        "stats": stats,
        "title_text": title_text,
        "page_setup": page_setup,
        "table_paragraphs": table_paragraph_count,
        "equation_paragraphs": equation_paragraph_count,
        "formatted_footnotes": formatted_footnote_count,
        "resized_images": resized_image_count,
        "cover_generated": cover_generated,
        "table_of_contents_inserted": toc_inserted,
        "format_options": resolved_format_options,
        "outline": outline,
    }


# ============================================================
# 封面+正文合并
# ============================================================
def merge_cover_and_body(cover_path: str, body_path: str, output_path: str, progress_callback=None, format_options=None):
    """
    合并封面文档和正文文档。

    封面保持原样不做排版处理，正文按学术论文规范排版，
    合并后正文部分页码从 1 开始。

    Returns:
        排版结果 dict（成功）或 False（失败）
    """
    import tempfile
    from contextlib import suppress

    cover_file = Path(cover_path)
    body_file = Path(body_path)

    if not cover_file.exists():
        logger.error(f"封面文件不存在：{cover_path}")
        return False

    if not body_file.exists():
        logger.error(f"正文文件不存在：{body_path}")
        return False

    formatted_body_path = None
    try:
        resolved_format_options = resolve_format_options(format_options)
        emit_progress(progress_callback, 1, "正在读取封面与正文文档")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            formatted_body_path = tmp.name

        body_format_options = resolved_format_options.copy()
        body_format_options["format_footnotes"] = False
        body_result = format_academic_paper(
            body_path,
            formatted_body_path,
            progress_callback=progress_callback,
            format_options=body_format_options,
        )
        if not body_result:
            return False

        from docxcompose.composer import Composer

        emit_progress(progress_callback, 4, "正在合并封面与排版后的正文")
        formatted_body_doc = Document(formatted_body_path)
        ensure_document_starts_with_page_break(formatted_body_doc)

        cover_doc = Document(cover_path)
        cover_section_count = len(cover_doc.sections)
        composer = Composer(cover_doc)
        composer.append(formatted_body_doc)

        # 在正文首节重启页码为 1
        merged_sections = list(cover_doc.sections)
        if len(merged_sections) > cover_section_count:
            body_first_section = merged_sections[cover_section_count]
            sect_pr = body_first_section._sectPr
            pg_num_type = sect_pr.find(qn("w:pgNumType"))
            if pg_num_type is None:
                pg_num_type = OxmlElement("w:pgNumType")
                sect_pr.append(pg_num_type)
            pg_num_type.set(qn("w:start"), "1")

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        emit_progress(progress_callback, 4, "正在写入合并后的输出文档")
        composer.save(output_path)
        formatted_footnote_count = 0
        if resolved_format_options["format_footnotes"]:
            emit_progress(progress_callback, 4, "正在统一合并文档中的脚注格式")
            formatted_footnote_count = format_docx_footnotes(output_path)
        if isinstance(body_result, dict):
            body_result["formatted_footnotes"] = formatted_footnote_count
            body_result["format_options"] = resolved_format_options

        logger.info(f"合并完成！封面 + 排版正文 → {output_path}")
        return body_result

    except Exception as e:
        logger.error(f"合并文档失败: {e}", exc_info=True)
        return False
    finally:
        if formatted_body_path:
            with suppress(OSError):
                Path(formatted_body_path).unlink()


# ============================================================
# 命令行入口
# ============================================================
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python format_paper.py <输入文件.docx> [输出文件.docx]")
        print("示例：python format_paper.py 论文初稿.docx 论文_排版后.docx")
        sys.exit(1)

    input_doc = sys.argv[1]

    if len(sys.argv) >= 3:
        output_doc = sys.argv[2]
    else:
        # 默认输出文件名：在原文件名后添加 "_formatted" 后缀
        p = Path(input_doc)
        output_doc = str(p.parent / f"{p.stem}_formatted{p.suffix}")

    success = format_academic_paper(input_doc, output_doc)
    sys.exit(0 if success else 1)
