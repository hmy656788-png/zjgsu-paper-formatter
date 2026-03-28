import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document

import app as app_module


TEST_DOC_PATH = Path(__file__).resolve().parent.parent / "test_input.docx"


class AppRoutesTestCase(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        base_dir = Path(self.temp_dir.name)

        self.upload_dir = base_dir / "uploads"
        self.output_dir = base_dir / "outputs"
        self.upload_dir.mkdir()
        self.output_dir.mkdir()

        self.original_upload_folder = app_module.UPLOAD_FOLDER
        self.original_output_folder = app_module.OUTPUT_FOLDER
        app_module.UPLOAD_FOLDER = self.upload_dir
        app_module.OUTPUT_FOLDER = self.output_dir
        app_module.PROGRESS_JOBS.clear()

        self.client = app_module.app.test_client()
        self.test_doc_bytes = TEST_DOC_PATH.read_bytes()

    def tearDown(self):
        app_module.UPLOAD_FOLDER = self.original_upload_folder
        app_module.OUTPUT_FOLDER = self.original_output_folder
        app_module.PROGRESS_JOBS.clear()
        self.temp_dir.cleanup()

    def test_format_preserves_unicode_filename(self):
        response = self.client.post(
            "/api/format",
            data={"file": (BytesIO(self.test_doc_bytes), "我的论文终稿.docx")},
            content_type="multipart/form-data",
        )

        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertEqual(data["original_name"], "我的论文终稿")
        self.assertEqual(data["download_name"], "我的论文终稿_排版后.docx")

    def test_format_generates_cover_from_form_fields(self):
        response = self.client.post(
            "/api/format",
            data={
                "file": (BytesIO(self.test_doc_bytes), "test_input.docx"),
                "generate_cover": "1",
                "cover_title": "《大数据挖掘》期末大作业",
                "college": "工商管理学院",
                "teacher": "刘璇",
                "class_name": "国商2301",
                "student_name": "何旻洋",
                "student_id": "2320100731",
            },
            content_type="multipart/form-data",
        )

        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertTrue(data["success"])
        self.assertTrue(data["format_summary"]["cover_generated"])

        output_doc = Document(next(self.output_dir.glob("*_output.docx")))
        self.assertEqual(output_doc.paragraphs[2].text, "《大数据挖掘》期末大作业")
        self.assertEqual(output_doc.tables[0].cell(0, 1).text, "工商管理学院")
        self.assertEqual(output_doc.tables[0].cell(4, 1).text, "2320100731")

    def test_format_cleans_up_uploaded_temp_file(self):
        response = self.client.post(
            "/api/format",
            data={"file": (BytesIO(self.test_doc_bytes), "test_input.docx")},
            content_type="multipart/form-data",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(list(self.upload_dir.glob("*.docx")), [])
        self.assertEqual(len(list(self.output_dir.glob("*_output.docx"))), 1)

    def test_download_rejects_non_generated_filename(self):
        rogue_output = self.output_dir / "manual.docx"
        rogue_output.write_bytes(b"placeholder")

        response = self.client.get("/api/download/manual.docx")

        self.assertEqual(response.status_code, 404)
        self.assertEqual(
            response.get_json(),
            {"success": False, "error": "文件不存在或已过期"},
        )

    def test_download_sanitizes_custom_download_name(self):
        generated_output = self.output_dir / "deadbeef_output.docx"
        generated_output.write_bytes(b"placeholder")

        response = self.client.get(
            "/api/download/deadbeef_output.docx?name=../../final-version"
        )
        try:
            self.assertEqual(response.status_code, 200)
            content_disposition = response.headers.get("Content-Disposition", "")
            self.assertIn("final-version.docx", content_disposition)
            self.assertNotIn("..", content_disposition)
            self.assertEqual(response.headers.get("Cache-Control"), "no-store")
            self.assertEqual(response.headers.get("X-Content-Type-Options"), "nosniff")
        finally:
            response.close()

    def test_format_text_rejects_invalid_json_body(self):
        response = self.client.post(
            "/api/format_text",
            data="not-json",
            headers={"Content-Type": "application/json"},
        )

        self.assertEqual(response.status_code, 400)
        self.assertEqual(
            response.get_json(),
            {"success": False, "error": "请输入有效的文字内容"},
        )

    def test_api_not_found_returns_json(self):
        response = self.client.get("/api/not-found")

        self.assertEqual(response.status_code, 404)
        self.assertEqual(
            response.get_json(),
            {"success": False, "error": "接口不存在，请确认请求地址是否正确。"},
        )
        self.assertEqual(response.headers.get("Cache-Control"), "no-store")
        self.assertEqual(response.headers.get("X-Content-Type-Options"), "nosniff")

    def test_api_method_not_allowed_returns_json(self):
        response = self.client.get("/api/format")

        self.assertEqual(response.status_code, 405)
        self.assertEqual(
            response.get_json(),
            {"success": False, "error": "请求方法不受支持，请检查接口调用方式。"},
        )
        self.assertEqual(response.headers.get("Cache-Control"), "no-store")
        self.assertEqual(response.headers.get("X-Content-Type-Options"), "nosniff")

    def test_api_health_reports_ready_storage(self):
        response = self.client.get("/api/health")

        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertTrue(data["success"])
        self.assertEqual(data["status"], "ok")
        self.assertTrue(data["storage"]["upload_ready"])
        self.assertTrue(data["storage"]["output_ready"])
        self.assertEqual(response.headers.get("Cache-Control"), "no-store")
        self.assertEqual(response.headers.get("X-Content-Type-Options"), "nosniff")

    def test_api_internal_error_does_not_leak_exception_details(self):
        original_formatter = app_module.format_academic_paper

        def broken_formatter(*_args, **_kwargs):
            raise RuntimeError("secret stack detail")

        app_module.format_academic_paper = broken_formatter

        try:
            response = self.client.post(
                "/api/format",
                data={"file": (BytesIO(self.test_doc_bytes), "test_input.docx")},
                content_type="multipart/form-data",
            )
        finally:
            app_module.format_academic_paper = original_formatter

        self.assertEqual(response.status_code, 500)
        self.assertEqual(
            response.get_json(),
            {"success": False, "error": "服务器内部错误，请稍后重试。"},
        )
        self.assertEqual(response.headers.get("Cache-Control"), "no-store")
        self.assertEqual(response.headers.get("X-Content-Type-Options"), "nosniff")

    def test_format_text_returns_preview_summary(self):
        response = self.client.post(
            "/api/format_text",
            json={
                "text": (
                    "基于多元回归模型的城市化研究\n"
                    "摘要：这是摘要内容\n"
                    "关键词：城市化 回归\n"
                    "1 引言\n"
                    "1.1 研究背景\n"
                    "1.1.1 研究假设\n"
                    "正文内容\n"
                    "参考文献\n"
                    "[1] 张三. 学术论文写作规范[J]. 高教研究, 2024."
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertTrue(data["success"])
        self.assertEqual(data["format_summary"]["page_setup"]["page_size"], "A4")
        self.assertEqual(data["format_summary"]["stats"]["heading_l3"], 1)
        self.assertEqual(data["format_summary"]["stats"]["reference_entry"], 1)
        self.assertEqual(len(data["preview"]["highlights"]), 4)
        self.assertTrue(any(item["level"] == "h3" for item in data["preview"]["outline"]))
        self.assertEqual(response.headers.get("Cache-Control"), "no-store")
        self.assertEqual(response.headers.get("X-Content-Type-Options"), "nosniff")

    def test_format_text_preview_preserves_section_heading_after_references(self):
        response = self.client.post(
            "/api/format_text",
            json={
                "text": (
                    "基于多元回归模型的城市化研究\n"
                    "摘要：这是摘要内容\n"
                    "关键词：城市化 回归\n"
                    "参考文献\n"
                    "[1] 张三. 学术论文写作规范[J]. 高教研究, 2024.\n"
                    "致谢\n"
                    "感谢导师和同学们的帮助。"
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertEqual(data["format_summary"]["stats"]["reference_entry"], 1)
        self.assertEqual(data["format_summary"]["stats"]["section_heading"], 1)
        self.assertTrue(any(item["level"] == "section" for item in data["preview"]["outline"]))

    def test_format_async_streams_progress_and_exposes_result(self):
        start_response = self.client.post(
            "/api/format_async",
            data={"file": (BytesIO(self.test_doc_bytes), "test_input.docx")},
            content_type="multipart/form-data",
        )

        self.assertEqual(start_response.status_code, 202)
        start_data = start_response.get_json()
        self.assertTrue(start_data["success"])
        self.assertIn("/api/jobs/", start_data["events_url"])
        self.assertIn("/api/jobs/", start_data["result_url"])

        events_response = self.client.get(start_data["events_url"], buffered=True)
        events_text = events_response.get_data(as_text=True)
        self.assertEqual(events_response.status_code, 200)
        self.assertIn("event: progress", events_text)
        self.assertIn("event: complete", events_text)
        self.assertIn("正在生成输出文档", events_text)

        result_response = self.client.get(start_data["result_url"])
        result_data = result_response.get_json()
        self.assertEqual(result_response.status_code, 200)
        self.assertTrue(result_data["success"])
        self.assertEqual(result_data["format_summary"]["page_setup"]["page_size"], "A4")
        self.assertIn("preview", result_data)

    def test_format_text_async_streams_progress_and_exposes_result(self):
        start_response = self.client.post(
            "/api/format_text_async",
            json={
                "text": (
                    "基于多元回归模型的城市化研究\n"
                    "摘要：这是摘要内容\n"
                    "关键词：城市化 回归\n"
                    "1 引言\n"
                    "1.1 研究背景\n"
                    "正文内容"
                )
            },
        )

        self.assertEqual(start_response.status_code, 202)
        start_data = start_response.get_json()
        self.assertTrue(start_data["success"])

        events_response = self.client.get(start_data["events_url"], buffered=True)
        events_text = events_response.get_data(as_text=True)
        self.assertEqual(events_response.status_code, 200)
        self.assertIn("event: progress", events_text)
        self.assertIn("event: complete", events_text)
        self.assertIn("识别标题层级", events_text)

        result_response = self.client.get(start_data["result_url"])
        result_data = result_response.get_json()
        self.assertEqual(result_response.status_code, 200)
        self.assertTrue(result_data["success"])
        self.assertEqual(result_data["format_summary"]["stats"]["heading_l2"], 1)


if __name__ == "__main__":
    unittest.main()
