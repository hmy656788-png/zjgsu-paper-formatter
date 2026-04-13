import base64
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.shared import Inches
from lxml import etree

from format_paper import (
    apply_document_layout,
    classify_paragraph,
    ensure_document_ends_with_page_break,
    ensure_document_starts_with_page_break,
    find_title_paragraph_index,
    format_academic_paper,
    format_academic_paper_from_text,
    generate_cover_page,
    merge_cover_and_body,
    ParagraphType,
    split_text_to_paragraphs,
    _find_field_paragraph_indices,
)

try:
    from docxcompose.composer import Composer as _DocxComposer  # noqa: F401
    HAS_DOCXCOMPOSE = True
except ImportError:  # pragma: no cover - optional dependency fallback
    HAS_DOCXCOMPOSE = False


class FormatPaperFromTextTestCase(unittest.TestCase):
    CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
    PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
    FOOTNOTE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
    WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    @staticmethod
    def assert_run_uses_mixed_font_pair(test_case, run):
        r_fonts = run._element.rPr.rFonts
        test_case.assertEqual(r_fonts.get(qn("w:eastAsia")), "宋体")
        test_case.assertEqual(r_fonts.get(qn("w:ascii")), "Times New Roman")
        test_case.assertEqual(r_fonts.get(qn("w:hAnsi")), "Times New Roman")

    def assert_paragraph_has_numbering(self, paragraph, ilvl: int) -> int:
        num_pr = paragraph._element.pPr.find(qn("w:numPr"))
        self.assertIsNotNone(num_pr)
        self.assertEqual(num_pr.find(qn("w:ilvl")).get(qn("w:val")), str(ilvl))
        return int(num_pr.find(qn("w:numId")).get(qn("w:val")))

    def assert_paragraph_has_no_numbering(self, paragraph):
        p_pr = paragraph._element.pPr
        num_pr = None if p_pr is None else p_pr.find(qn("w:numPr"))
        self.assertIsNone(num_pr)

    def assert_numbering_overrides(self, doc, num_id: int, expected_starts: dict[int, int]):
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        num = numbering.num_having_numId(num_id)

        for ilvl, start_value in expected_starts.items():
            lvl_override = next(
                override
                for override in num.findall("./" + qn("w:lvlOverride"))
                if override.get(qn("w:ilvl")) == str(ilvl)
            )
            start_override = lvl_override.find(qn("w:startOverride"))
            self.assertIsNotNone(start_override)
            self.assertEqual(start_override.get(qn("w:val")), str(start_value))

    def inject_simple_footnote(self, docx_path: Path, footnote_text: str):
        with ZipFile(docx_path, "r") as source:
            entries = source.infolist()
            payloads = {entry.filename: source.read(entry.filename) for entry in entries}

        content_types = etree.fromstring(payloads["[Content_Types].xml"])
        override_tag = f"{{{self.CONTENT_TYPES_NS}}}Override"
        if not any(node.get("PartName") == "/word/footnotes.xml" for node in content_types.findall(override_tag)):
            override = etree.Element(override_tag)
            override.set("PartName", "/word/footnotes.xml")
            override.set(
                "ContentType",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
            )
            content_types.append(override)
        payloads["[Content_Types].xml"] = etree.tostring(
            content_types,
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )

        rels = etree.fromstring(payloads["word/_rels/document.xml.rels"])
        relationship_tag = f"{{{self.PACKAGE_REL_NS}}}Relationship"
        if not any(node.get("Type") == self.FOOTNOTE_REL_TYPE for node in rels.findall(relationship_tag)):
            relationship = etree.Element(relationship_tag)
            relationship.set("Id", "rIdFootnotes")
            relationship.set("Type", self.FOOTNOTE_REL_TYPE)
            relationship.set("Target", "footnotes.xml")
            rels.append(relationship)
        payloads["word/_rels/document.xml.rels"] = etree.tostring(
            rels,
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )

        document_xml = etree.fromstring(payloads["word/document.xml"])
        body = document_xml.find(qn("w:body"))
        target_paragraph = body.findall(qn("w:p"))[-1]
        target_paragraph.append(
            parse_xml(
                f'<w:r {nsdecls("w")}>'
                '<w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
                '<w:footnoteReference w:id="2"/>'
                "</w:r>"
            )
        )
        payloads["word/document.xml"] = etree.tostring(
            document_xml,
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )

        payloads["word/footnotes.xml"] = (
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:footnotes {nsdecls("w")}>'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
            '<w:footnote w:id="2"><w:p>'
            '<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r>'
            f"<w:r><w:t xml:space=\"preserve\"> {footnote_text}</w:t></w:r>"
            "</w:p></w:footnote>"
            "</w:footnotes>"
        ).encode("utf-8")

        with ZipFile(docx_path, "w") as target:
            written = set()
            for entry in entries:
                target.writestr(entry, payloads[entry.filename])
                written.add(entry.filename)
            for name, data in payloads.items():
                if name not in written:
                    target.writestr(name, data)

    def remove_style_definition(self, docx_path: Path, style_id: str):
        with ZipFile(docx_path, "r") as source:
            entries = source.infolist()
            payloads = {entry.filename: source.read(entry.filename) for entry in entries}

        styles_xml = payloads.get("word/styles.xml")
        self.assertIsNotNone(styles_xml)

        styles_root = etree.fromstring(styles_xml)
        style_tag = f"{{{self.WORD_NS}}}style"
        name_tag = f"{{{self.WORD_NS}}}name"
        removed = False
        for style in list(styles_root.findall(style_tag)):
            current_style_id = style.get(qn("w:styleId"))
            name_node = style.find(name_tag)
            current_name = name_node.get(qn("w:val")) if name_node is not None else ""
            if current_style_id == style_id or current_name == style_id:
                styles_root.remove(style)
                removed = True

        self.assertTrue(removed)
        payloads["word/styles.xml"] = etree.tostring(
            styles_root,
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )

        with ZipFile(docx_path, "w") as target:
            written = set()
            for entry in entries:
                target.writestr(entry, payloads[entry.filename])
                written.add(entry.filename)
            for name, data in payloads.items():
                if name not in written:
                    target.writestr(name, data)

    def test_split_text_to_paragraphs_normalizes_line_endings(self):
        text = "第一段\r\n\r\n第二段\r第三段\n"

        self.assertEqual(
            split_text_to_paragraphs(text),
            ["第一段", "", "第二段", "第三段"],
        )

    def test_format_academic_paper_from_text_preserves_blank_lines_without_newline_chars(self):
        text = "第一段\r\n\r\n第二段\r\n"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            self.assertTrue(format_academic_paper_from_text(text, str(output_path)))

            doc = Document(str(output_path))
            self.assertEqual([paragraph.text for paragraph in doc.paragraphs], ["第一段", "", "第二段"])
        finally:
            output_path.unlink(missing_ok=True)

    def test_find_title_paragraph_index_only_when_abstract_follows(self):
        doc = Document()
        doc.add_paragraph("基于多元回归模型的城市化研究")
        doc.add_paragraph("作者：张三")
        doc.add_paragraph("摘要：这是摘要内容")

        self.assertEqual(find_title_paragraph_index(doc.paragraphs), 0)

    def test_find_title_paragraph_index_skips_metadata_prefix(self):
        doc = Document()
        doc.add_paragraph("作者：张三")
        doc.add_paragraph("基于多元回归模型的城市化研究")
        doc.add_paragraph("摘要：这是摘要内容")

        self.assertEqual(find_title_paragraph_index(doc.paragraphs), 1)

    def test_find_title_paragraph_index_rejects_metadata_without_real_title(self):
        doc = Document()
        doc.add_paragraph("作者：张三")
        doc.add_paragraph("摘要：这是摘要内容")
        doc.add_paragraph("关键词：城市化 回归")

        self.assertIsNone(find_title_paragraph_index(doc.paragraphs))

    def test_find_title_paragraph_index_accepts_english_abstract_heading(self):
        doc = Document()
        doc.add_paragraph("跨境电商场景下供应链韧性研究")
        doc.add_paragraph("Abstract")
        doc.add_paragraph("This paper studies supply chain resilience.")

        self.assertEqual(find_title_paragraph_index(doc.paragraphs), 0)

    def test_format_academic_paper_from_text_formats_detected_title(self):
        text = "基于多元回归模型的城市化研究\n摘要：这是摘要内容\n关键词：城市化 回归\n1 引言\n正文内容"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            self.assertTrue(format_academic_paper_from_text(text, str(output_path)))

            doc = Document(str(output_path))
            title = doc.paragraphs[0]
            self.assertEqual(title.text, "基于多元回归模型的城市化研究")
            self.assertEqual(title.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(title.runs[0].font.bold)
            self.assertEqual(title.runs[0].font.size.pt, 18.0)
        finally:
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_from_text_generates_cover_from_cover_info(self):
        text = (
            "企业数字化转型对绿色技术创新的影响研究\n"
            "摘要：这是摘要内容\n"
            "关键词：数字化 创新\n"
            "正文内容"
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            summary = format_academic_paper_from_text(
                text,
                str(output_path),
                cover_info={
                    "cover_title": "《大数据挖掘》期末大作业",
                    "college": "工商管理学院",
                    "teacher": "刘璇",
                    "class_name": "国商2301",
                    "student_name": "何旻洋",
                    "student_id": "2320100731",
                },
            )

            self.assertTrue(summary["cover_generated"])
            output_doc = Document(str(output_path))
            self.assertEqual(output_doc.paragraphs[2].text, "《大数据挖掘》期末大作业")
            self.assertEqual(output_doc.tables[0].cell(0, 1).text, "工商管理学院")
            self.assertEqual(output_doc.tables[0].cell(4, 1).text, "2320100731")
        finally:
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_from_text_applies_page_layout_and_page_number(self):
        text = "基于多元回归模型的城市化研究\n摘要：这是摘要内容\n关键词：城市化 回归\n1 引言\n正文内容"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            summary = format_academic_paper_from_text(text, str(output_path))

            self.assertIsInstance(summary, dict)
            self.assertEqual(summary["page_setup"]["page_size"], "A4")
            self.assertEqual(summary["page_setup"]["header_text"], "基于多元回归模型的城市化研究")

            doc = Document(str(output_path))
            section = doc.sections[0]

            self.assertAlmostEqual(section.page_width.cm, 21.0, places=1)
            self.assertAlmostEqual(section.page_height.cm, 29.7, places=1)
            self.assertAlmostEqual(section.top_margin.cm, 2.54, places=1)
            self.assertAlmostEqual(section.left_margin.cm, 3.18, places=1)
            self.assertEqual(section.header.paragraphs[0].text, "基于多元回归模型的城市化研究")
            self.assertIn("PAGE", section.footer._element.xml)
        finally:
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_from_text_formats_heading_l3_and_references(self):
        text = (
            "基于多元回归模型的城市化研究\n"
            "摘要：这是摘要内容\n"
            "关键词：城市化 回归\n"
            "1 引言\n"
            "1.1 研究背景\n"
            "1.1.1 研究假设\n"
            "正文内容\n"
            "参考文献\n"
            "[1] 张三. 学术论文写作规范[J]. 高教研究, 2024.\n"
            "[2] 李四. 文献整理方法[M]. 北京: 科学出版社, 2023."
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            summary = format_academic_paper_from_text(text, str(output_path))

            self.assertIsInstance(summary, dict)
            self.assertEqual(summary["stats"]["heading_l3"], 1)
            self.assertEqual(summary["stats"]["references_heading"], 1)
            self.assertEqual(summary["stats"]["reference_entry"], 2)
            self.assertTrue(any(item["level"] == "h3" for item in summary["outline"]))
            self.assertTrue(any(item["level"] == "references" for item in summary["outline"]))

            doc = Document(str(output_path))
            heading_l1 = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "1 引言")
            heading_l2 = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "1.1 研究背景")
            heading_l3 = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "1.1.1 研究假设")
            references_heading = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "参考文献：")
            reference_entry = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("[1] 张三."))
            toc_heading = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "目录")
            toc_field = next(
                paragraph for paragraph in doc.paragraphs
                if 'TOC \\o "1-3" \\h \\z \\u' in paragraph._element.xml
            )
            toc_page_break = next(
                paragraph for paragraph in doc.paragraphs
                if 'w:type="page"' in paragraph._element.xml
            )

            self.assertEqual(heading_l3.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertTrue(heading_l3.runs[0].font.bold)
            self.assertEqual(heading_l3.runs[0].font.size.pt, 12.0)
            self.assertEqual(heading_l3.paragraph_format.first_line_indent.pt, 0.0)
            self.assertEqual(heading_l3._element.pPr.find(qn("w:outlineLvl")).get(qn("w:val")), "2")
            self.assert_paragraph_has_no_numbering(heading_l1)
            self.assert_paragraph_has_no_numbering(heading_l2)
            self.assert_paragraph_has_no_numbering(heading_l3)

            self.assertEqual(references_heading.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(references_heading.text, "参考文献：")
            self.assertEqual(references_heading.runs[0].font.size.pt, 12.0)
            self.assertEqual(reference_entry.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertAlmostEqual(reference_entry.paragraph_format.left_indent.pt, 0.0, places=1)
            self.assertAlmostEqual(reference_entry.paragraph_format.first_line_indent.pt, 21.0, places=1)
            self.assertEqual(reference_entry.paragraph_format.line_spacing, 1.0)
            self.assertEqual(reference_entry.runs[0].font.size.pt, 10.0)
            self.assertEqual(toc_heading.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertIn('TOC \\o "1-3" \\h \\z \\u', toc_field._element.xml)
            self.assertIn('w:type="page"', toc_page_break._element.xml)
            self.assertIn("w:updateFields", doc.settings.element.xml)
        finally:
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_from_text_stops_references_before_acknowledgements(self):
        text = (
            "基于多元回归模型的城市化研究\n"
            "摘要：这是摘要内容\n"
            "关键词：城市化 回归\n"
            "参考文献\n"
            "[1] 张三. 学术论文写作规范[J]. 高教研究, 2024.\n"
            "致谢\n"
            "感谢导师和同学们的帮助。"
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            summary = format_academic_paper_from_text(text, str(output_path))

            self.assertIsInstance(summary, dict)
            self.assertEqual(summary["stats"]["reference_entry"], 1)
            self.assertEqual(summary["stats"]["section_heading"], 1)
            self.assertTrue(any(item["level"] == "section" for item in summary["outline"]))

            doc = Document(str(output_path))
            acknowledgements = doc.paragraphs[5]
            acknowledgement_body = doc.paragraphs[6]

            self.assertEqual(acknowledgements.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(acknowledgement_body.paragraph_format.first_line_indent.pt, 24.0)
        finally:
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_from_text_shortens_running_header_for_long_title(self):
        text = (
            "基于多源异构数据融合与深度学习模型的中国城市高质量发展测度及影响机制研究\n"
            "摘要：这是摘要内容\n"
            "关键词：城市化 回归\n"
            "正文内容"
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            summary = format_academic_paper_from_text(text, str(output_path))

            self.assertIsInstance(summary, dict)
            self.assertTrue(summary["page_setup"]["header_text"].endswith("..."))
            self.assertLessEqual(len(summary["page_setup"]["header_text"]), 28)

            doc = Document(str(output_path))
            self.assertEqual(doc.sections[0].header.paragraphs[0].text, summary["page_setup"]["header_text"])
        finally:
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_from_text_formats_english_abstract_keywords(self):
        text = (
            "跨境电商场景下供应链韧性研究\n"
            "Abstract\n"
            "This paper studies supply chain resilience under cross-border e-commerce settings.\n"
            "Keywords: supply chain resilience; cross-border e-commerce\n"
            "1 Introduction\n"
            "正文内容"
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            summary = format_academic_paper_from_text(text, str(output_path))

            self.assertIsInstance(summary, dict)
            self.assertEqual(summary["stats"]["english_abstract_heading"], 1)
            self.assertEqual(summary["stats"]["english_abstract"], 1)
            self.assertEqual(summary["stats"]["english_keywords"], 1)

            doc = Document(str(output_path))
            abstract_heading = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "Abstract")
            abstract_body = next(
                paragraph for paragraph in doc.paragraphs
                if paragraph.text.startswith("This paper studies supply chain resilience")
            )
            keywords = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("Keywords:"))
            introduction = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "Introduction")

            self.assertEqual(abstract_heading.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(abstract_heading.runs[0].font.bold)
            self.assertEqual(abstract_heading.runs[0].font.size.pt, 12.0)

            self.assertEqual(abstract_body.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertEqual(abstract_body.paragraph_format.first_line_indent.pt, 0.0)
            self.assertEqual(abstract_body.runs[0].font.size.pt, 12.0)
            self.assertEqual(abstract_body._element.pPr.find(qn("w:widowControl")).get(qn("w:val")), "true")

            self.assertEqual(keywords.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertTrue(keywords.runs[0].font.bold)
            self.assertEqual(keywords.text, "Keywords: supply chain resilience; cross-border e-commerce")
            self.assertEqual(keywords._element.pPr.find(qn("w:widowControl")).get(qn("w:val")), "true")
            self.assert_paragraph_has_numbering(introduction, 0)
        finally:
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_from_text_preserves_explicit_heading_number_offsets(self):
        text = (
            "数字化转型场景下企业韧性研究\n"
            "摘要：这是摘要内容\n"
            "关键词：数字化 韧性\n"
            "3 研究设计\n"
            "3.2 数据来源\n"
            "3.2.4 稳健性检验\n"
            "正文内容"
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            summary = format_academic_paper_from_text(text, str(output_path))

            self.assertIsInstance(summary, dict)

            doc = Document(str(output_path))
            heading_l1 = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "3 研究设计")
            heading_l2 = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "3.2 数据来源")
            heading_l3 = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "3.2.4 稳健性检验")

            self.assert_paragraph_has_no_numbering(heading_l1)
            self.assert_paragraph_has_no_numbering(heading_l2)
            self.assert_paragraph_has_no_numbering(heading_l3)
        finally:
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_infers_numbering_from_heading_styles(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            input_path = temp_path / "style_heading_input.docx"
            output_path = temp_path / "style_heading_output.docx"

            doc = Document()
            doc.add_paragraph("平台治理视角下数字化协同研究")
            doc.add_paragraph("摘要：这是摘要内容")
            doc.add_paragraph("关键词：平台治理 协同")
            doc.add_paragraph("引言", style="Heading 1")
            doc.add_paragraph("研究背景", style="Heading 2")
            doc.add_paragraph("研究假设", style="Heading 3")
            doc.add_paragraph("正文内容")
            doc.save(str(input_path))

            summary = format_academic_paper(str(input_path), str(output_path))

            self.assertEqual(summary["stats"]["heading_l1"], 1)
            self.assertEqual(summary["stats"]["heading_l2"], 1)
            self.assertEqual(summary["stats"]["heading_l3"], 1)

            output_doc = Document(str(output_path))
            heading_l1 = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "引言")
            heading_l2 = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "研究背景")
            heading_l3 = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "研究假设")

            heading_l1_num_id = self.assert_paragraph_has_numbering(heading_l1, 0)
            heading_l2_num_id = self.assert_paragraph_has_numbering(heading_l2, 1)
            heading_l3_num_id = self.assert_paragraph_has_numbering(heading_l3, 2)

            self.assert_numbering_overrides(output_doc, heading_l1_num_id, {0: 1})
            self.assert_numbering_overrides(output_doc, heading_l2_num_id, {0: 1, 1: 1})
            self.assert_numbering_overrides(output_doc, heading_l3_num_id, {0: 1, 1: 1, 2: 1})

    def test_format_academic_paper_does_not_auto_number_plain_short_body(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            input_path = temp_path / "plain_short_body_input.docx"
            output_path = temp_path / "plain_short_body_output.docx"

            doc = Document()
            doc.add_paragraph("短标题误判保护测试")
            doc.add_paragraph("摘要：这是摘要内容")
            doc.add_paragraph("关键词：误判 测试")
            doc.add_paragraph("研究意义")
            doc.add_paragraph("这里是正文展开内容。")
            doc.save(str(input_path))

            summary = format_academic_paper(str(input_path), str(output_path))

            self.assertEqual(summary["stats"]["heading_l1"], 0)
            self.assertEqual(summary["stats"]["heading_l2"], 0)
            self.assertEqual(summary["stats"]["heading_l3"], 0)

            output_doc = Document(str(output_path))
            paragraph = next(item for item in output_doc.paragraphs if item.text == "研究意义")
            num_pr = paragraph._element.pPr.find(qn("w:numPr")) if paragraph._element.pPr is not None else None
            self.assertIsNone(num_pr)
            self.assertEqual(paragraph.paragraph_format.first_line_indent.pt, 24.0)

    def test_format_academic_paper_preserves_equation_paragraphs(self):
        omml = parse_xml(
            r"""
            <m:oMathPara %s>
              <m:oMath>
                <m:r>
                  <m:t>x=1</m:t>
                </m:r>
              </m:oMath>
            </m:oMathPara>
            """ % nsdecls("m")
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            input_path = temp_path / "equation_input.docx"
            output_path = temp_path / "equation_output.docx"

            doc = Document()
            doc.add_paragraph("含公式的论文标题")
            doc.add_paragraph("摘要：这是摘要内容")
            doc.add_paragraph("关键词：公式 测试")
            equation_paragraph = doc.add_paragraph()
            equation_paragraph._element.append(omml)
            doc.save(str(input_path))

            summary = format_academic_paper(str(input_path), str(output_path))

            self.assertEqual(summary["equation_paragraphs"], 1)

            output_doc = Document(str(output_path))
            preserved_equation = output_doc.paragraphs[3]
            self.assertIn("oMath", preserved_equation._element.xml)
            self.assertEqual(preserved_equation.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(preserved_equation.paragraph_format.first_line_indent.pt, 0.0)

    def test_format_academic_paper_unifies_footnote_fonts_and_sizes(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            input_path = temp_path / "footnote_input.docx"
            output_path = temp_path / "footnote_output.docx"

            doc = Document()
            doc.add_paragraph("含脚注的论文标题")
            doc.add_paragraph("摘要：这是摘要内容")
            doc.add_paragraph("关键词：脚注 测试")
            doc.add_paragraph("正文里有一个脚注引用")
            doc.save(str(input_path))
            self.inject_simple_footnote(input_path, "脚注内容 Footnote 123")

            summary = format_academic_paper(str(input_path), str(output_path))

            self.assertEqual(summary["formatted_footnotes"], 1)

            with ZipFile(output_path, "r") as archive:
                footnotes_xml = archive.read("word/footnotes.xml")

            footnotes_root = etree.fromstring(footnotes_xml)
            footnote = next(
                node
                for node in footnotes_root.findall(qn("w:footnote"))
                if node.get(qn("w:id")) == "2"
            )
            footnote_paragraph = footnote.find(qn("w:p"))
            runs = footnote.findall(".//" + qn("w:r"))
            reference_rpr = runs[0].find(qn("w:rPr"))
            content_rpr = runs[1].find(qn("w:rPr"))
            footnote_ppr = footnote_paragraph.find(qn("w:pPr"))
            footnote_spacing = footnote_ppr.find(qn("w:spacing"))
            footnote_indent = footnote_ppr.find(qn("w:ind"))
            footnote_jc = footnote_ppr.find(qn("w:jc"))

            self.assertEqual(reference_rpr.find(qn("w:rStyle")).get(qn("w:val")), "FootnoteReference")
            self.assertEqual(reference_rpr.find(qn("w:vertAlign")).get(qn("w:val")), "superscript")
            self.assertEqual(reference_rpr.find(qn("w:sz")).get(qn("w:val")), "20")
            self.assertEqual(content_rpr.find(qn("w:rFonts")).get(qn("w:eastAsia")), "宋体")
            self.assertEqual(content_rpr.find(qn("w:rFonts")).get(qn("w:ascii")), "Times New Roman")
            self.assertEqual(content_rpr.find(qn("w:sz")).get(qn("w:val")), "20")
            self.assertEqual(footnote_spacing.get(qn("w:before")), "0")
            self.assertEqual(footnote_spacing.get(qn("w:after")), "0")
            self.assertEqual(footnote_spacing.get(qn("w:line")), "240")
            self.assertEqual(footnote_spacing.get(qn("w:lineRule")), "auto")
            self.assertEqual(footnote_indent.get(qn("w:left")), "0")
            self.assertEqual(footnote_indent.get(qn("w:right")), "0")
            self.assertEqual(footnote_indent.get(qn("w:firstLine")), "0")
            self.assertEqual(footnote_jc.get(qn("w:val")), "left")
            self.assertEqual(footnote_ppr.find(qn("w:widowControl")).get(qn("w:val")), "true")

    def test_format_academic_paper_handles_complex_doc_with_images_table_and_lists(self):
        tiny_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9VE3d2wAAAAASUVORK5CYII="
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            image_path = temp_path / "pixel.png"
            input_path = temp_path / "complex_input.docx"
            output_path = temp_path / "complex_output.docx"
            image_path.write_bytes(tiny_png)

            doc = Document()
            doc.add_paragraph("跨境电商场景下 Python 模型驱动的供应链韧性研究")
            doc.add_paragraph("摘要：本文使用 Python 3.11 和 Cross-border E-commerce 数据进行分析。")
            doc.add_paragraph("关键词：Python Cross-border E-commerce 2024")

            heading = doc.add_paragraph()
            heading.add_run("1")
            heading.add_run().add_break()
            heading.add_run("引言")

            doc.add_paragraph("本文基于 Python 3.11 与 Cross-border E-commerce 2024 数据进行实证分析。")
            doc.add_paragraph("First bullet uses Python 3.11", style="List Bullet")

            picture_1 = doc.add_paragraph()
            picture_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture_1.add_run().add_picture(str(image_path), width=Inches(0.2))
            doc.add_paragraph("【图9】 系统架构图")

            picture_2 = doc.add_paragraph()
            picture_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture_2.add_run().add_picture(str(image_path), width=Inches(0.2))
            doc.add_paragraph("图8 模型流程图")

            doc.add_paragraph("表88 样本描述统计")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "变量"
            table.cell(0, 1).text = "Python 3.11"
            table.cell(1, 0).text = "平台"
            table.cell(1, 1).text = "Cross-border E-commerce 2024"
            doc.add_paragraph("注：样本区间为 2020-2024 年。")
            doc.add_paragraph("来源：Python 爬取与企业年报整理。")

            doc.add_paragraph("参考文献")
            doc.add_paragraph("[9] Smith, John. Python-based trade analytics[J]. 2024.")
            doc.save(str(input_path))

            summary = format_academic_paper(str(input_path), str(output_path))
            self.assertIsInstance(summary, dict)
            self.assertEqual(summary["stats"]["figure_caption"], 2)
            self.assertEqual(summary["stats"]["table_caption"], 1)
            self.assertEqual(summary["stats"]["caption_note"], 2)
            self.assertEqual(summary["table_paragraphs"], 4)

            output_doc = Document(str(output_path))
            self.assertEqual(len(output_doc.inline_shapes), 2)

            self.assertEqual(output_doc.paragraphs[3].text, "1 引言")
            self.assertEqual(output_doc.paragraphs[7].text, "图 1 系统架构图")
            self.assertEqual(output_doc.paragraphs[9].text, "图 2 模型流程图")
            self.assertEqual(output_doc.paragraphs[10].text, "表 1 样本描述统计")
            self.assertEqual(output_doc.paragraphs[11].text, "注：样本区间为 2020-2024 年。")
            self.assertEqual(output_doc.paragraphs[12].text, "来源：Python 爬取与企业年报整理。")
            self.assertEqual(output_doc.paragraphs[5].style.name, "List Bullet")
            self.assertEqual(output_doc.paragraphs[6].paragraph_format.first_line_indent.pt, 0.0)
            self.assert_paragraph_has_no_numbering(output_doc.paragraphs[3])
            self.assertEqual(output_doc.paragraphs[3]._element.pPr.find(qn("w:keepNext")).get(qn("w:val")), "true")
            self.assertEqual(output_doc.paragraphs[3]._element.pPr.find(qn("w:keepLines")).get(qn("w:val")), "true")
            self.assertEqual(output_doc.paragraphs[6]._element.pPr.find(qn("w:keepNext")).get(qn("w:val")), "true")
            self.assertEqual(output_doc.paragraphs[7]._element.pPr.find(qn("w:keepLines")).get(qn("w:val")), "true")
            self.assertIsNone(output_doc.paragraphs[7]._element.pPr.find(qn("w:keepNext")))
            self.assertEqual(output_doc.paragraphs[10]._element.pPr.find(qn("w:keepNext")).get(qn("w:val")), "true")
            self.assertEqual(output_doc.paragraphs[4]._element.pPr.find(qn("w:widowControl")).get(qn("w:val")), "true")
            self.assertEqual(output_doc.paragraphs[11]._element.pPr.find(qn("w:widowControl")).get(qn("w:val")), "true")
            self.assertEqual(output_doc.paragraphs[14]._element.pPr.find(qn("w:widowControl")).get(qn("w:val")), "true")

            body_run = output_doc.paragraphs[4].runs[0]
            table_run = output_doc.tables[0].cell(1, 1).paragraphs[0].runs[0]
            self.assert_run_uses_mixed_font_pair(self, body_run)
            self.assert_run_uses_mixed_font_pair(self, table_run)

            caption_note = output_doc.paragraphs[11]
            source_note = output_doc.paragraphs[12]
            reference_entry = output_doc.paragraphs[14]
            self.assertEqual(caption_note.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertAlmostEqual(caption_note.paragraph_format.first_line_indent.pt, 0.0, places=1)
            self.assertEqual(caption_note.paragraph_format.line_spacing, 1.0)
            self.assertTrue(caption_note.runs[0].font.bold)
            self.assertEqual(caption_note.runs[0].font.size.pt, 10.5)
            self.assertEqual(source_note.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertTrue(source_note.runs[0].font.bold)
            self.assertEqual(source_note.runs[0].font.size.pt, 10.5)
            self.assertAlmostEqual(reference_entry.paragraph_format.left_indent.pt, 0.0, places=1)
            self.assertAlmostEqual(reference_entry.paragraph_format.first_line_indent.pt, 21.0, places=1)
            self.assertEqual(reference_entry.paragraph_format.line_spacing, 1.0)
            self.assertEqual(reference_entry.runs[0].font.size.pt, 10.0)

            tbl_borders = output_doc.tables[0]._tbl.tblPr.find(qn("w:tblBorders"))
            top_left_borders = output_doc.tables[0].cell(0, 0)._tc.tcPr.find(qn("w:tcBorders"))
            bottom_left_borders = output_doc.tables[0].cell(1, 0)._tc.tcPr.find(qn("w:tcBorders"))
            first_row_tr_pr = output_doc.tables[0].rows[0]._tr.trPr
            second_row_tr_pr = output_doc.tables[0].rows[1]._tr.trPr
            header_paragraph = output_doc.tables[0].cell(0, 0).paragraphs[0]
            self.assertIsNone(tbl_borders)
            self.assertEqual(output_doc.tables[0].alignment, WD_TABLE_ALIGNMENT.CENTER)
            self.assertEqual(output_doc.tables[0].cell(0, 0).vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
            self.assertEqual(top_left_borders.find(qn("w:top")).get(qn("w:val")), "single")
            self.assertEqual(top_left_borders.find(qn("w:bottom")).get(qn("w:val")), "single")
            self.assertEqual(top_left_borders.find(qn("w:left")).get(qn("w:val")), "none")
            self.assertEqual(top_left_borders.find(qn("w:right")).get(qn("w:val")), "none")
            self.assertEqual(bottom_left_borders.find(qn("w:top")).get(qn("w:val")), "none")
            self.assertEqual(bottom_left_borders.find(qn("w:bottom")).get(qn("w:val")), "single")
            self.assertEqual(header_paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(header_paragraph.runs[0].font.bold)
            self.assertEqual(first_row_tr_pr.find(qn("w:tblHeader")).get(qn("w:val")), "true")
            self.assertEqual(first_row_tr_pr.find(qn("w:cantSplit")).get(qn("w:val")), "true")
            self.assertEqual(second_row_tr_pr.find(qn("w:cantSplit")).get(qn("w:val")), "true")

    def test_format_academic_paper_scales_oversized_images_to_page_width(self):
        tiny_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9VE3d2wAAAAASUVORK5CYII="
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            image_path = temp_path / "oversized.png"
            input_path = temp_path / "oversized_input.docx"
            output_path = temp_path / "oversized_output.docx"
            image_path.write_bytes(tiny_png)

            doc = Document()
            doc.add_paragraph("超宽图片版式测试")
            doc.add_paragraph("摘要：这是摘要内容。")
            doc.add_paragraph("关键词：图片 缩放 测试")

            picture = doc.add_paragraph()
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture.add_run().add_picture(str(image_path), width=Inches(8))
            doc.add_paragraph("图 1 超宽测试图")
            doc.save(str(input_path))

            summary = format_academic_paper(str(input_path), str(output_path))
            self.assertIsInstance(summary, dict)
            self.assertEqual(summary["resized_images"], 1)

            output_doc = Document(str(output_path))
            self.assertEqual(len(output_doc.inline_shapes), 1)

            printable_width = (
                int(output_doc.sections[0].page_width)
                - int(output_doc.sections[0].left_margin)
                - int(output_doc.sections[0].right_margin)
            )
            resized_shape = output_doc.inline_shapes[0]

            self.assertLessEqual(int(resized_shape.width), printable_width)
            self.assertEqual(int(resized_shape.width), int(resized_shape.height))
            self.assertLess(int(resized_shape.width), int(Inches(8)))

    def test_format_academic_paper_does_not_misclassify_body_explanations_as_caption_notes(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            input_path = temp_path / "body_explanation_input.docx"
            output_path = temp_path / "body_explanation_output.docx"

            doc = Document()
            doc.add_paragraph("平台治理视角下数字化协同研究")
            doc.add_paragraph("摘要：这是摘要内容。")
            doc.add_paragraph("关键词：平台治理 协同")
            doc.add_paragraph("说明：这是正文中的说明句，不应被当作图表附注。")
            doc.save(str(input_path))

            summary = format_academic_paper(str(input_path), str(output_path))

            self.assertEqual(summary["stats"]["caption_note"], 0)

            output_doc = Document(str(output_path))
            body_paragraph = output_doc.paragraphs[3]
            self.assertAlmostEqual(body_paragraph.paragraph_format.first_line_indent.pt, 24.0, places=1)
            self.assertEqual(body_paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertEqual(body_paragraph.runs[0].font.size.pt, 12.0)
            self.assertFalse(body_paragraph.runs[0].font.bold)

    def test_generate_cover_page_inserts_cover_table_and_page_break(self):
        info_dict = {
            "title": "企业数字化转型对绿色技术创新的影响研究",
            "cover_title": "《大数据挖掘》期末大作业",
            "college": "工商管理学院",
            "teacher": "刘璇",
            "class_name": "国商2301",
            "student_name": "何旻洋",
            "student_id": "2320100731",
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            doc = Document()
            doc.add_paragraph("这是已经完成正文排版的第一页内容。")
            doc.add_paragraph("这是正文第二段。")
            apply_document_layout(doc, "这是正文运行页眉")

            self.assertTrue(generate_cover_page(doc, info_dict))
            doc.save(str(output_path))

            output_doc = Document(str(output_path))

            self.assertEqual(output_doc.paragraphs[0].text, "")
            self.assertEqual(output_doc.paragraphs[1].text, "")
            self.assertEqual(output_doc.paragraphs[2].text, info_dict["cover_title"])
            self.assertEqual(output_doc.paragraphs[3].text, "")
            self.assertEqual(output_doc.paragraphs[5].text, "这是已经完成正文排版的第一页内容。")
            self.assertIn('w:type="page"', output_doc.paragraphs[4]._element.xml)
            self.assertEqual(len(output_doc.inline_shapes), 2)

            title_run = output_doc.paragraphs[2].runs[0]
            self.assertEqual(output_doc.paragraphs[2].paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(title_run.font.bold)
            self.assertEqual(title_run.font.size.pt, 26.0)

            cover_table = output_doc.tables[0]
            self.assertEqual(len(cover_table.rows), 5)
            self.assertEqual(cover_table.cell(0, 0).text, "学院")
            self.assertEqual(cover_table.cell(0, 1).text, "工商管理学院")
            self.assertEqual(cover_table.cell(4, 1).text, "2320100731")

            label_borders = cover_table.cell(0, 0)._tc.tcPr.find(qn("w:tcBorders"))
            value_borders = cover_table.cell(0, 1)._tc.tcPr.find(qn("w:tcBorders"))
            self.assertEqual(label_borders.find(qn("w:bottom")).get(qn("w:val")), "none")
            self.assertEqual(value_borders.find(qn("w:top")).get(qn("w:val")), "none")
            self.assertEqual(value_borders.find(qn("w:left")).get(qn("w:val")), "none")
            self.assertEqual(value_borders.find(qn("w:right")).get(qn("w:val")), "none")
            self.assertEqual(value_borders.find(qn("w:bottom")).get(qn("w:val")), "single")

            paragraph_borders = cover_table.cell(0, 1).paragraphs[0]._element.pPr.find(qn("w:pBdr"))
            self.assertIsNotNone(paragraph_borders)
            self.assertEqual(paragraph_borders.find(qn("w:bottom")).get(qn("w:val")), "single")

            label_run = cover_table.cell(0, 0).paragraphs[0].runs[0]
            value_run = cover_table.cell(4, 1).paragraphs[0].runs[0]
            self.assert_run_uses_mixed_font_pair(self, label_run)
            self.assert_run_uses_mixed_font_pair(self, value_run)

            section = output_doc.sections[0]
            self.assertTrue(section.different_first_page_header_footer)
            self.assertEqual(section.first_page_header.paragraphs[0].text, "")
            self.assertEqual(section.first_page_footer.paragraphs[0].text, "")
            self.assertEqual(section.header.paragraphs[0].text, "这是正文运行页眉")
        finally:
            output_path.unlink(missing_ok=True)

    def test_generate_cover_page_skips_when_title_is_missing(self):
        doc = Document()
        doc.add_paragraph("正文内容")

        self.assertFalse(generate_cover_page(doc, {"student_name": "何旻洋"}))
        self.assertEqual([paragraph.text for paragraph in doc.paragraphs], ["正文内容"])

    def test_ensure_document_ends_with_page_break_appends_break(self):
        doc = Document()
        doc.add_paragraph("封面内容")

        ensure_document_ends_with_page_break(doc)

        self.assertTrue(doc.paragraphs)
        self.assertIn('w:type="page"', doc.paragraphs[-1]._element.xml)

    def test_ensure_document_starts_with_page_break_prepends_break(self):
        doc = Document()
        doc.add_paragraph("正文标题")

        ensure_document_starts_with_page_break(doc)

        self.assertTrue(doc.paragraphs)
        self.assertIn('w:type="page"', doc.paragraphs[0]._element.xml)
        self.assertEqual(doc.paragraphs[1].text, "正文标题")

    @unittest.skipUnless(HAS_DOCXCOMPOSE, "docxcompose 未安装，跳过封面合并测试")
    def test_merge_cover_and_body_preserves_cover_content(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            cover_path = temp_path / "cover.docx"
            body_path = temp_path / "body.docx"
            output_path = temp_path / "merged.docx"

            cover_doc = Document()
            cover_title = cover_doc.add_paragraph("浙江工商大学课程论文")
            cover_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_title.add_run("封面副标题").bold = True
            cover_doc.add_paragraph("学生姓名：张三")
            cover_doc.save(str(cover_path))

            body_doc = Document()
            body_doc.add_paragraph("平台治理视角下数字化协同研究")
            body_doc.add_paragraph("摘要：这是摘要内容。")
            body_doc.add_paragraph("关键词：平台治理 协同")
            body_doc.add_paragraph("1 引言")
            body_doc.add_paragraph("这是正文第一页。")
            body_doc.save(str(body_path))

            summary = merge_cover_and_body(str(cover_path), str(body_path), str(output_path))

            self.assertIsInstance(summary, dict)

            merged_doc = Document(str(output_path))
            self.assertEqual(merged_doc.paragraphs[0].text, "浙江工商大学课程论文封面副标题")
            self.assertEqual(merged_doc.paragraphs[1].text, "学生姓名：张三")
            self.assertNotIn('w:type="page"', merged_doc.paragraphs[0]._element.xml)
            self.assertNotIn('w:type="page"', merged_doc.paragraphs[1]._element.xml)
            self.assertIn('w:type="page"', merged_doc.paragraphs[2]._element.xml)
            self.assertEqual(merged_doc.paragraphs[3].text, "平台治理视角下数字化协同研究")
            self.assertEqual(merged_doc.paragraphs[0].paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(merged_doc.paragraphs[0].runs[-1].bold)

            original_cover_doc = Document(str(cover_path))
            self.assertEqual(original_cover_doc.paragraphs[0].text, "浙江工商大学课程论文封面副标题")
            self.assertEqual(original_cover_doc.paragraphs[1].text, "学生姓名：张三")
            self.assertEqual(len(original_cover_doc.paragraphs), 2)


    # ------------------------------------------------------------------
    # 标题正则扩展测试
    # ------------------------------------------------------------------

    def test_classify_heading_with_special_characters(self):
        """标题中包含引号、书名号、百分号等特殊字符应正确识别。"""
        cases = [
            ("1 基于\u201c双碳\u201d目标的分析", ParagraphType.HEADING_L1),
            ("2 《国富论》的核心观点", ParagraphType.HEADING_L1),
            ("3 GDP增长率达到8.5%的原因", ParagraphType.HEADING_L1),
            ("1.1 \u201c互联网+\u201d背景下的研究", ParagraphType.HEADING_L2),
            ("2.3 基于Black\u2013Scholes模型", ParagraphType.HEADING_L2),
            ("1.1.1 \u201c双碳\u201d与ESG", ParagraphType.HEADING_L3),
            ("2.1.1 变量\u00b7指标选取", ParagraphType.HEADING_L3),
        ]
        for text, expected in cases:
            with self.subTest(text=text):
                self.assertEqual(classify_paragraph(text), expected)

    def test_classify_heading_l1_without_space_after_period(self):
        """\"1.引言\" 形式（句点后无空格）应识别为一级标题。"""
        self.assertEqual(classify_paragraph("1.引言"), ParagraphType.HEADING_L1)
        self.assertEqual(classify_paragraph("2.研究设计"), ParagraphType.HEADING_L1)
        # "1引言"（无任何分隔）应保持为正文，避免误判"1年"之类
        self.assertEqual(classify_paragraph("1引言"), ParagraphType.BODY)

    def test_classify_rejects_date_and_counter_patterns_as_headings(self):
        """正文中 \"数字 空格 X 空格 数字\" 的日期/计数表达不应被识别为标题。"""
        cases = [
            "1 月 1 日开始实验",
            "1 第 2 章开始",
            "2024 年 3 月的数据",
            "2 第 3 条规定",
        ]
        for text in cases:
            with self.subTest(text=text):
                self.assertEqual(classify_paragraph(text), ParagraphType.BODY)

    def test_classify_toc_entries_as_body(self):
        """手动输入的目录条目（含点引导+页码）不应被识别为标题或图表标题。"""
        cases = [
            "1 引言........3",
            "1.1 研究背景......5",
            "2 研究设计............7",
            "图 1 数据分布图............3",
            "表 2 变量定义……5",
        ]
        for text in cases:
            with self.subTest(text=text):
                self.assertEqual(classify_paragraph(text), ParagraphType.BODY)

    def test_caption_parses_chapter_prefixed_index(self):
        """章节前缀的图表编号（如 图1-1、表2.3）应被完整识别为 index。"""
        from format_paper import match_caption_in_normalized_text, rebuild_caption_text

        cases = [
            ("图 1-1 变量分布", "1-1", "变量分布"),
            ("图1-1 变量分布", "1-1", "变量分布"),
            ("图 1.1 变量分布", "1.1", "变量分布"),
            ("表 2-3 数据统计", "2-3", "数据统计"),
            ("图 1.1.1 xxx", "1.1.1", "xxx"),
        ]
        for text, expected_index, expected_caption in cases:
            with self.subTest(text=text):
                result = match_caption_in_normalized_text(text)
                self.assertIsNotNone(result)
                _, match, _ = result
                self.assertEqual(match.group("index"), expected_index)
                self.assertEqual(match.group("caption"), expected_caption)
                # 重建标题时应使用自动编号，不应把旧 index 残留在 caption 中
                rebuilt = rebuild_caption_text(result[0], 5, match)
                self.assertNotIn("1-1", rebuilt)
                self.assertNotIn("2-3", rebuilt)

    def test_classify_section_heading_variants(self):
        """\"致谢\"、\"附录\" 等章节标题应容忍内部空格与附加编号。"""
        cases = [
            "致谢", "致 谢", "致  谢",
            "附录", "附 录", "附录A", "附录 A", "附录一",
            "作者简介", "作 者 简 介",
            "基金项目",
        ]
        for text in cases:
            with self.subTest(text=text):
                self.assertEqual(classify_paragraph(text), ParagraphType.SECTION_HEADING)

    # ------------------------------------------------------------------
    # 域段落（图目录/表目录/TOC）跳过测试
    # ------------------------------------------------------------------

    def test_find_field_paragraph_indices_detects_toc_style(self):
        """带 TOC 样式的段落应被识别为域段落。"""
        from docx.oxml import OxmlElement

        doc = Document()
        doc.add_paragraph("正文段落")
        p_toc = doc.add_paragraph("1 引言……3")
        # 模拟 TOC 样式：直接设置样式名
        p_toc.style = doc.styles.add_style("TOC 1", 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH

        paragraphs = list(doc.paragraphs)
        indices = _find_field_paragraph_indices(paragraphs)

        self.assertNotIn(0, indices)
        self.assertIn(1, indices)

    def test_find_field_paragraph_indices_detects_field_block(self):
        """被 fldChar begin/end 包裹的段落应被识别为域段落。"""
        from docx.oxml import OxmlElement

        doc = Document()
        doc.add_paragraph("正文")

        # 创建含有 fldChar begin 的段落
        p_field_start = doc.add_paragraph()
        run_begin = p_field_start.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run_begin._element.append(fld_begin)
        run_instr = p_field_start.add_run()
        instr = OxmlElement("w:instrText")
        instr.text = 'TOC \\o "1-3"'
        run_instr._element.append(instr)
        run_sep = p_field_start.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run_sep._element.append(fld_sep)

        # 域内容段落（模拟目录条目）
        doc.add_paragraph("图 1 某某图……5")

        # 域结束段落
        p_field_end = doc.add_paragraph()
        run_end = p_field_end.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end._element.append(fld_end)

        paragraphs = list(doc.paragraphs)
        indices = _find_field_paragraph_indices(paragraphs)

        # 域开始段落（含 begin 和 separate）不算域内容，但域中间的内容段落应被标记
        self.assertNotIn(0, indices)  # 正文段落
        # p_field_start 包含 begin，之后 field_depth > 0，所以它自身也在域内
        self.assertIn(1, indices)
        self.assertIn(2, indices)  # 域内容段落

    def test_format_academic_paper_skips_toc_style_paragraphs(self):
        """带 TOC 样式的段落不应被当作图表标题重新格式化。"""
        doc = Document()
        doc.add_paragraph("论文标题")
        doc.add_paragraph("摘要：这是摘要")
        doc.add_paragraph("关键词：测试")
        doc.add_paragraph("1 引言")
        doc.add_paragraph("正文段落")

        # 添加一个模拟图目录条目
        toc_entry = doc.add_paragraph("图 1 数据分布图")
        toc_entry.style = doc.styles.add_style("Table of Figures", 1)

        doc.add_paragraph("图 1 数据分布图")  # 真正的图表标题

        with tempfile.NamedTemporaryFile(suffix=".docx") as f:
            doc.save(f.name)
            summary = format_academic_paper(f.name, f.name)

        # 只应识别到 1 个图标题（真正的），而非 2 个
        self.assertEqual(summary["stats"]["figure_caption"], 1)

    def test_format_academic_paper_skips_toc_hyperlink_entries(self):
        """指向 _Toc 书签的自动目录条目不应被当作真实图标题。"""
        from docx.oxml import OxmlElement

        doc = Document()
        doc.add_paragraph("论文标题")
        doc.add_paragraph("摘要：这是摘要")
        doc.add_paragraph("关键词：测试")
        doc.add_paragraph("1 引言")
        doc.add_paragraph("正文段落")

        toc_entry = doc.add_paragraph()
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), "_Toc123456789")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "图 1 数据分布图……5"
        run.append(text)
        hyperlink.append(run)
        toc_entry._p.append(hyperlink)

        doc.add_paragraph("图 1 数据分布图")

        with tempfile.NamedTemporaryFile(suffix=".docx") as f:
            doc.save(f.name)
            summary = format_academic_paper(f.name, f.name)

        self.assertEqual(summary["stats"]["figure_caption"], 1)

    def test_format_academic_paper_handles_documents_without_normal_style(self):
        """缺少 Normal 样式的模板文档也应能完成排版。"""
        doc = Document()
        doc.add_paragraph("A Strategic Analysis of Luckin Coffee")
        doc.add_paragraph("Abstract")
        doc.add_paragraph("Keywords: coffee strategy")
        doc.add_paragraph("1. Introduction")
        doc.add_paragraph("正文内容")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            input_path = Path(handle.name)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            doc.save(input_path)
            self.remove_style_definition(input_path, "Normal")

            summary = format_academic_paper(str(input_path), str(output_path))

            self.assertIsInstance(summary, dict)
            self.assertTrue(output_path.exists())
            self.assertEqual(summary["page_setup"]["page_size"], "A4")
        finally:
            input_path.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_infers_styled_english_step_heading(self):
        """英文 Heading 2 中包含内部数字时，仍应按标题而非正文处理。"""
        doc = Document()
        doc.add_paragraph("Test Report")
        doc.add_paragraph("Abstract")
        heading = doc.add_paragraph("3.1 Step 1: Mission")
        heading.style = "Heading 2"
        doc.add_paragraph("The mission statement is the conceptual point of departure.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            input_path = Path(handle.name)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            doc.save(input_path)
            summary = format_academic_paper(str(input_path), str(output_path))

            self.assertEqual(summary["stats"]["heading_l2"], 1)

            output_doc = Document(str(output_path))
            heading_paragraph = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "Step 1: Mission")
            self.assertEqual(heading_paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertEqual(heading_paragraph.paragraph_format.first_line_indent.pt, 0.0)
        finally:
            input_path.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_left_aligns_english_body_and_h1(self):
        """纯英文正文与一级标题应避免两端拉伸和居中错位。"""
        doc = Document()
        heading = doc.add_paragraph("3. Phase One: Goal Setting")
        heading.style = "Heading 1"
        body = doc.add_paragraph("The mission statement is the conceptual point of departure for any strategic plan.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            input_path = Path(handle.name)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            doc.save(input_path)
            summary = format_academic_paper(str(input_path), str(output_path))

            self.assertEqual(summary["stats"]["heading_l1"], 1)

            output_doc = Document(str(output_path))
            heading_paragraph = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "Phase One: Goal Setting")
            body_paragraph = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text.startswith("The mission statement"))

            self.assertEqual(heading_paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertEqual(body_paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertEqual(body_paragraph.paragraph_format.first_line_indent.pt, 24.0)
        finally:
            input_path.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_formats_long_english_table_caption_and_references(self):
        """较长的英文表题与英文参考文献段落都应按对应版式格式化。"""
        doc = Document()
        doc.add_paragraph("Abstract")
        doc.add_paragraph("Abstract: This report evaluates Luckin Coffee's strategic development.")
        doc.add_paragraph("Keywords: coffee strategy")
        doc.add_paragraph("Table 1. Luckin Coffee Selected Performance Indicators, 2019–2024")
        doc.add_paragraph("Source: Luckin Coffee SEC filings (Forms 6-K and annual reports), 2019–2024.")
        doc.add_paragraph("References")
        doc.add_paragraph("Drucker, P. F. (1973). Management: Tasks, responsibilities, practices. Harper & Row.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            input_path = Path(handle.name)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            doc.save(input_path)
            summary = format_academic_paper(str(input_path), str(output_path))

            self.assertEqual(summary["stats"]["table_caption"], 1)
            self.assertEqual(summary["stats"]["caption_note"], 1)
            self.assertEqual(summary["stats"]["references_heading"], 1)
            self.assertEqual(summary["stats"]["reference_entry"], 1)

            output_doc = Document(str(output_path))
            caption_paragraph = next(
                paragraph
                for paragraph in output_doc.paragraphs
                if paragraph.text.startswith("Table 1 ")
            )
            source_paragraph = next(
                paragraph
                for paragraph in output_doc.paragraphs
                if paragraph.text.startswith("Source:")
            )
            references_heading = next(
                paragraph
                for paragraph in output_doc.paragraphs
                if paragraph.text == "References"
            )
            reference_entry = next(
                paragraph
                for paragraph in output_doc.paragraphs
                if paragraph.text.startswith("Drucker, P. F. (1973).")
            )

            self.assertEqual(caption_paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertAlmostEqual(caption_paragraph.runs[0].font.size.pt, 10.5, places=1)
            self.assertEqual(source_paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertAlmostEqual(source_paragraph.runs[0].font.size.pt, 10.5, places=1)
            self.assertTrue(source_paragraph.runs[0].bold)
            self.assertEqual(references_heading.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(reference_entry.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertAlmostEqual(reference_entry.paragraph_format.first_line_indent.pt, 21.0, places=1)
        finally:
            input_path.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)

    def test_format_academic_paper_matches_english_template_layout(self):
        """英文整篇模板应保留封面区、编号标题、拆行表题与参考文献的样式。"""
        doc = Document()
        for text in [
            "A Strategic Analysis of Luckin Coffee Through the",
            "Marketing Planning Framework:",
            "From Accounting Scandal to Market Leadership",
        ]:
            paragraph = doc.add_paragraph(text)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        author = doc.add_paragraph("He Minyang")
        author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        abstract_heading = doc.add_paragraph("Abstract")
        abstract_heading.style = "Heading 1"
        doc.add_paragraph("This report applies McDonald's marketing planning framework to analyse Luckin Coffee.")
        doc.add_paragraph("Keywords: Luckin Coffee, marketing planning framework, competitive strategy")

        intro = doc.add_paragraph("1. Introduction")
        intro.style = "Heading 1"
        doc.add_paragraph("The Chinese freshly brewed coffee market has undergone a structural transformation.")
        doc.add_paragraph("Table 1")
        doc.add_paragraph("Luckin Coffee Selected Performance Indicators, 2019–2024")
        doc.add_paragraph("Note. Luckin Coffee SEC filings (Forms 6-K and annual reports), 2019–2024.")

        step = doc.add_paragraph("2.1 Step 1: Mission")
        step.style = "Heading 2"
        substep = doc.add_paragraph("2.1.1 PESTLE Analysis")
        substep.style = "Heading 3"

        references = doc.add_paragraph("References")
        references.style = "Heading 1"
        doc.add_paragraph("Aaker, D. A. (1996). Building strong brands. Free Press.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            input_path = Path(handle.name)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            output_path = Path(handle.name)

        try:
            doc.save(input_path)
            summary = format_academic_paper(str(input_path), str(output_path))

            self.assertEqual(summary["stats"]["table_caption"], 1)
            self.assertEqual(summary["stats"]["caption_note"], 1)
            self.assertEqual(summary["stats"]["references_heading"], 1)
            self.assertEqual(summary["stats"]["reference_entry"], 1)

            output_doc = Document(str(output_path))
            title_line = next(
                paragraph
                for paragraph in output_doc.paragraphs
                if paragraph.text == "A Strategic Analysis of Luckin Coffee Through the"
            )
            author_line = next(
                paragraph
                for paragraph in output_doc.paragraphs
                if paragraph.text == "He Minyang"
            )
            abstract_heading = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "Abstract")
            keywords = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text.startswith("Keywords:"))
            intro = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "1. Introduction")
            table_label = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "Table 1")
            table_title = next(
                paragraph for paragraph in output_doc.paragraphs
                if paragraph.text == "Luckin Coffee Selected Performance Indicators, 2019–2024"
            )
            note = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text.startswith("Note."))
            step = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "2.1 Step 1: Mission")
            substep = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "2.1.1 PESTLE Analysis")
            references_heading = next(paragraph for paragraph in output_doc.paragraphs if paragraph.text == "References")
            reference_entry = next(
                paragraph for paragraph in output_doc.paragraphs
                if paragraph.text.startswith("Aaker, D. A. (1996).")
            )

            self.assertEqual(title_line.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(title_line.runs[0].bold)
            self.assertEqual(author_line.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertFalse(author_line.runs[0].bold)

            self.assertEqual(abstract_heading.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertAlmostEqual(abstract_heading.runs[0].font.size.pt, 14.0, places=1)
            self.assertAlmostEqual(keywords.paragraph_format.first_line_indent.pt, 36.0, places=1)
            self.assertFalse(keywords.runs[0].bold)

            self.assertEqual(intro.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(intro.text, "1. Introduction")
            self.assertAlmostEqual(intro.runs[0].font.size.pt, 14.0, places=1)
            self.assertEqual(step.text, "2.1 Step 1: Mission")
            self.assertAlmostEqual(step.runs[0].font.size.pt, 13.0, places=1)
            self.assertEqual(substep.text, "2.1.1 PESTLE Analysis")

            self.assertEqual(table_label.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertTrue(table_label.runs[0].bold)
            self.assertAlmostEqual(table_label.runs[0].font.size.pt, 12.0, places=1)
            self.assertEqual(table_title.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertFalse(table_title.runs[0].bold)
            self.assertAlmostEqual(note.runs[0].font.size.pt, 10.0, places=1)
            self.assertFalse(note.runs[0].bold)

            self.assertEqual(references_heading.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertAlmostEqual(references_heading.runs[0].font.size.pt, 14.0, places=1)
            self.assertAlmostEqual(reference_entry.paragraph_format.first_line_indent.pt, -36.0, places=1)
            self.assertAlmostEqual(reference_entry.paragraph_format.left_indent.pt, 36.0, places=1)
            self.assertAlmostEqual(reference_entry.runs[0].font.size.pt, 11.0, places=1)
        finally:
            input_path.unlink(missing_ok=True)
            output_path.unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()
